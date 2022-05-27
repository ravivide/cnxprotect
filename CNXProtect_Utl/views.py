import keyword
from urllib import response

from dill import objects
from django.core.files.storage import FileSystemStorage
import pandas as pd
from CNX_Protect import settings
from django.views import generic
# from .models import File
from .models import Master_Data
import os, datetime
from django.contrib import messages
from CNX_Protect.settings import BASE_DIR
from django.contrib.auth.models import auth
from django.http import HttpRequest, Http404
from django.shortcuts import render, redirect
from webpage.decorators import unauthenticated_user
from django.contrib.auth.decorators import login_required
# from .resources import DictionaryResource
from .models import Dictionary
# from tablib import Dataset
# Create your views here.

# from django.shortcuts import render_to_response
from django.template import RequestContext
from django.http import HttpResponseRedirect
# from django.core.urlresolvers import reverse

from CNXProtect_Utl.models import Document
from CNXProtect_Utl.forms import DocumentForm

from .models import Dictionary
from .models import Dictionary
from .models import Dictionary1
from CNXProtect_Utl.models import Add_Keywords
from CNXProtect_Utl.models import Customer
# Create your views here.
from django.contrib import messages
from django.contrib.auth.models import Group
from django.contrib.auth.models import auth
from django.shortcuts import render, redirect
from CNXProtect_Utl.models import Customer
from webpage.decorators import unauthenticated_user
# Create your views here.
from webpage.forms import CreateUserForm
from webpage.decorators import unauthenticated_user, allowed_users  # , admin_only
from .models import *
from CNX_Protect.settings import BASE_DIR
from .forms import UploadFileForm
from .models import File
from .models import File_1
from .models import File_pdf
# from ContractLby.views import cl_files
import requests


def index(request):
    return render(request, 'CNXProtect_Utl/index.html')


def blog(request):
    return render(request, 'CNXProtect_Utl/blog.html')


def single_work(request):
    return render(request, 'CNXProtect_Utl/single_work.html')


@login_required(login_url='login')
# @admin_only
def indexing_data5(request):
    customers = Customer.objects.all()
    return render(request, 'CNXProtect_Utl/indexing5.html')


def Contract_Library(request):
    return render(request, 'CNXProtect_Utl/Contract_Library.html')


def Contract_Library1(request):
    return render(request, 'CNXProtect_Utl/Contract_Library1.html')


def CLBT_reports(request):
    return render(request, 'CNXProtect_Utl/CLBT_reports.html')


def FMEA(request):
    return render(request, 'CNXProtect_Utl/FMEA.html')


def FMEA_Indexing(request):
    return render(request, 'CNXProtect_Utl/FMEA_Indexing.html')


def FMEA_Real(request):
    return render(request, 'CNXProtect_Utl/FMEA_Real.html')


def experiment1(request):
    return render(request, "CNXProtect_Utl/experiment1.html")


def experiment(request):
    return render(request, "CNXProtect_Utl/experiment.html")


def experiment2(request):
    if request.method == "POST":
        notesfile_1 = request.FILES['notesfile_1']
        p = notesfile_1
        # p1 = file.notesfile
        a = Master_Data(notesfile_1=notesfile_1)
        a.save()
    p3 = Master_Data.objects.values('notesfile_1')
    # print(Master_Data.objects.filter(notesfile_1__startwith = '0' ))
    p4 = pd.DataFrame(p3, columns=['notesfile_1'])
    print(p4)
    F_Len = len(p4.iloc[:]['notesfile_1'])
    print(F_Len)
    F_Count = p4.iloc[F_Len - 1]['notesfile_1']
    X_File = "C://Users/sumithra.r/Downloads/deployment/media/" + F_Count
    print(X_File)
    XL_File = pd.read_excel(X_File)
    # XL_File['Date_of_modified'] = pd.to_datetime(XL_File['Date_of_modified'])
    MSA_Name = request.GET.getlist('MSA_1')
    print("List of MSA", MSA_Name)
    Theme_Name = request.GET.getlist('Theme')
    # Time_in = request.GET['InputTime']
    # Time_out = request.GET['OutputTime']
    # mask = (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)
    # print(mask)
    # print(MSA_Name[0])
    # print(Theme_Name[0])
    # print(Time_in)
    # print(Time_out)

    DF17 = XL_File[XL_File['MSA'].isin(MSA_Name) & XL_File['Word'].isin(Theme_Name)]
    # DF17 = XL_File.loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]
    excel_filename_1 = "Contract_Library_Results_" + str(datetime.datetime.today().date()) + ".xlsx"
    DF17.to_excel(BASE_DIR + '/media/' + excel_filename_1)
    DF17_1 = '/media/' + excel_filename_1
    print(DF17)


# For Downloads
    if request.method == "POST":
        notesfile_1 = request.FILES['notesfile_1']
        p = notesfile_1
        # p1 = file.notesfile
        a = Master_Data(notesfile_1=notesfile_1)
        a.save()
    p3 = Master_Data.objects.values('notesfile_1')
    # print(Master_Data.objects.filter(notesfile_1__startwith = '0' ))
    p4 = pd.DataFrame(p3, columns=['notesfile_1'])
    print(p4)
    F_Len = len(p4.iloc[:]['notesfile_1'])
    print(F_Len)
    F_Count = p4.iloc[F_Len - 1]['notesfile_1']
    X_File = "C://Users/sumithra.r/Downloads/deployment/media/" + F_Count
    XL_File = pd.read_excel(X_File)
    excel_filename_3 = "Contract_Master_Data" + str(datetime.datetime.today().date()) + ".xlsx"
    XL_File.to_excel(BASE_DIR + '/media/' + excel_filename_3)
    df_filepath_Exe1 = '/media/' + excel_filename_3

    # ************************************************ FROM HERE FOR FMEA *****************************************************************
    if request.method == "POST":
        notesfile_11 = request.FILES['notesfile_11']
        p = notesfile_11
        # p1 = file.notesfile
        a = Master_Data1(notesfile_11=notesfile_11)
        a.save()
    p5 = Master_Data1.objects.values('notesfile_11')
    # print(Master_Data.objects.filter(notesfile_1__startwith = '0' ))
    p6 = pd.DataFrame(p5, columns=['notesfile_11'])
    print(p6)
    F_Len = len(p6.iloc[:]['notesfile_11'])
    print(F_Len)
    F_Count_1 = p6.iloc[F_Len - 1]['notesfile_11']
    X_File_1 = "C://Users/sumithra.r/Downloads/deployment/media/" + F_Count_1
    XL_File_1 = pd.read_excel(X_File_1)
    excel_filename_4 = "FMEA_Master_Data" + str(datetime.datetime.today().date()) + ".xlsx"
    XL_File_1.to_excel(BASE_DIR + '/media/' + excel_filename_4)
    df_filepath_Exe1_1 = '/media/' + excel_filename_4

    # print(pd.DataFrame(MSA_Name))
    # Account = XL_File['MSA'].unique()
    # DF11 = XL_File['MSA'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Theme/Topic'] == mask)]
    # #                                                                                   # (df['birth_date'] > start_date) & (df['birth_date'] <= end_date)
    # DF12 = XL_File['Word'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]
    # DF13 = XL_File['Sentence'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]
    # DF14 = XL_File['Theme/Topic'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]
    # DF15 = XL_File['File_Name/DirName'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]
    # # DF15df = pd.DataFrame(DF15)
    # DF16 = XL_File['Date_of_modified'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]
    # print(DF13)
    #
    # DF1_1 = DF11.to_string(index=False).rjust(10)
    # DF1_2 = DF12.to_string(index=False).rjust(10)
    # if DF13.empty:
    #     DF1_3 = 'No sentences have been extracted from the Account:', MSA_Name, 'and the Theme:', Theme_Name
    # else:
    #     DF1_3 = DF13.to_string(index=False).rjust(10)
    # DF1_4 = DF14.to_string(index=False).rjust(10)
    # DF1_5 = DF15.to_string(index=False).rjust(10)
    # DF1_6 = DF16.to_string(index=False).rjust(10)

    return render(request, 'CNXProtect_Utl/experiment2.html', {'DF17_1': DF17_1, 'df_filepath_Exe1':df_filepath_Exe1, 'df_filepath_Exe1_1':df_filepath_Exe1_1})
                  # {'Account': Account,
                       # 'DF1_1': DF1_1, 'DF1_2': DF1_2, 'DF1_3': DF1_3, 'DF1_4': DF1_4, 'DF1_5': DF1_5, 'DF1_6': DF1_6})

    # return render(request, 'CNXProtect_Utl/experiment.html')


def FMEA_Indexing(request):
    if request.method == "POST":
        notesfile_11 = request.FILES['notesfile_11']
        p = notesfile_11
        # p1 = file.notesfile
        # a = Master_Data1(notesfile_11=notesfile_11)
        a = Master_Data1(notesfile_11=notesfile_11)
        a.save()
    p3 = Master_Data1.objects.values('notesfile_11')
    # print(Master_Data.objects.filter(notesfile_1__startwith = '0' ))
    p4 = pd.DataFrame(p3, columns=['notesfile_11'])
    print(p4)
    F_Len = len(p4.iloc[:]['notesfile_11'])
    print(F_Len)
    F_Count = p4.iloc[F_Len - 1]['notesfile_11']
    X_File = "C://Users/sumithra.r/Downloads/deployment/media/" + F_Count
    print(X_File)
    XL_File = pd.read_excel(X_File)
    # XL_File['Date_of_modified'] = pd.to_datetime(XL_File['Date_of_modified'])
    MSA_Name = request.GET.getlist('MSA_1')
    print("List of MSA", MSA_Name)
    Theme_Name = request.GET.getlist('Theme')
    Geo_Name = request.GET.getlist('Geo')
    # Time_in = request.GET['InputTime']
    # Time_out = request.GET['OutputTime']
    # mask = (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)
    # print(mask)
    # print(MSA_Name[0])
    # print(Theme_Name[0])
    # print(Time_in)
    # print(Time_out)

    DF17 = XL_File[XL_File['MSA'].isin(MSA_Name) & XL_File['Key'].isin(Theme_Name) & XL_File['Geo'].isin(Geo_Name)]
    # DF17 = XL_File.loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]
    excel_filename_1 = "Indexing_report_" + str(datetime.datetime.today().date()) + ".xlsx"
    DF17.to_excel(BASE_DIR + '/media/' + excel_filename_1)
    DF17_1 = '/media/' + excel_filename_1
    print("The results are:", DF17)

    # print(pd.DataFrame(MSA_Name))
    # Account = XL_File['MSA'].unique()
    # DF11 = XL_File['MSA'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Theme/Topic'] == mask)]
    # #                                                                                   # (df['birth_date'] > start_date) & (df['birth_date'] <= end_date)
    # DF12 = XL_File['Word'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]
    # DF13 = XL_File['Sentence'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]
    # DF14 = XL_File['Theme/Topic'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]
    # DF15 = XL_File['File_Name/DirName'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]
    # # DF15df = pd.DataFrame(DF15)
    # DF16 = XL_File['Date_of_modified'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]
    # print(DF13)
    #
    # DF1_1 = DF11.to_string(index=False).rjust(10)
    # DF1_2 = DF12.to_string(index=False).rjust(10)
    # if DF13.empty:
    #     DF1_3 = 'No sentences have been extracted from the Account:', MSA_Name, 'and the Theme:', Theme_Name
    # else:
    #     DF1_3 = DF13.to_string(index=False).rjust(10)
    # DF1_4 = DF14.to_string(index=False).rjust(10)
    # DF1_5 = DF15.to_string(index=False).rjust(10)
    # DF1_6 = DF16.to_string(index=False).rjust(10)

    return render(request, 'CNXProtect_Utl/FMEA_Indexing.html',{'DF17_1':DF17_1})
                  # {'Account': Account,
                       # 'DF1_1': DF1_1, 'DF1_2': DF1_2, 'DF1_3': DF1_3, 'DF1_4': DF1_4, 'DF1_5': DF1_5, 'DF1_6': DF1_6})


def Indexing(request):
    return render(request, 'CNXProtect_Utl/Indexing.html')


@unauthenticated_user
def register(request):
    form = CreateUserForm()
    if request.method == 'POST':
        form = CreateUserForm(request.POST)
        if form.is_valid():
            user = form.save()
            username = form.cleaned_data.get('username')
            group = Group.objects.get_or_create(name='customer')
            user.groups.add(group)
            messages.success(request, 'Account was created for ' + username)
            return redirect('login')
    context = {'form': form}
    return render(request, 'CNXProtect_Utl/register.html', context)


@unauthenticated_user
def login(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']

        user = auth.authenticate(username=username, password=password)
        if user is not None:
            auth.login(request, user)
            return redirect('indexing_data5')
        else:
            messages.info(request, 'You are not authorized to use')
            return redirect('login')
    else:
        messages.info(request, '')
        return render(request, 'CNXProtect_Utl/login.html')


@login_required(login_url='login')
@allowed_users(allowed_roles=['admin', 'customer'])
def customer(request, pk_test):
    customer = Customer.objects.get(id=pk_test)
    orders = customer.order_set.all()
    order_count = orders.count()
    return render(request, 'CNXProtect_Utl/index.html', order_count)


def keyword_list(request):
    context = {'keyword_list': Dictionary.objects.all()}
    print(context)
    # print(keyword_list)
    return render(request, 'CNXProtect_Utl/Contract_Library1.html', context)


def insert_keyword_item(request: HttpRequest):
    Keyword = Dictionary(keywords=request.POST['keywords'])
    Keyword.save()
    Keys = Keyword.keywords
    print(Keys)
    # print(Dictionary.objects.all())
    return redirect('keyword_list')


def delete_keyword_item(request, keyword_id):
    keyword_to_delete = Dictionary.objects.get(id=keyword_id)
    keyword_to_delete.delete()
    return redirect('keyword_list')


def logout(request):
    auth.logout(request)
    return redirect('/')


def keyword_list1(request: HttpRequest):
    context = {'keyword_list1': Dictionary1.objects.all()}
    print(context)
    # print(keyword_list)
    return render(request, 'CNXProtect_Utl/FMEA_Real.html', context)


def insert_keyword_item1(request: HttpRequest):
    Keyword = Dictionary1(keywords=request.POST['keywords'])
    Keyword.save()
    Keys = Keyword.keywords
    print(Keys)
    # print(Dictionary.objects.all())
    return redirect('FM_Real')


def delete_keyword_item1(request, keyword_id):
    keyword_to_delete = Dictionary1.objects.get(id=keyword_id)
    keyword_to_delete.delete()
    return redirect('FM_Real')


def pdfword(request):
    import os
    import re
    import docx
    # import numpy
    import PyPDF4
    FileCount = 1
    flg = "lower1"
    print("*** Keyword Search - Using Dictionary ***")

    dec = re.compile('\d+\.\d+')
    dig = re.compile('\d+')
    alphabets = "([A-Za-z])"
    prefixes = "(Mr|St|Mrs|Ms|Dr)[.]"
    suffixes = "(Inc|Ltd|Jr|Sr|Co)"
    starters = "(Mr|Mrs|Ms|Dr|He\s|She\s|It\s|They\s|Their\s|Our\s|We\s|But\s|However\s|That\s|This\s|Wherever)"
    acronyms = "([A-Z][.][A-Z][.](?:[A-Z][.])?)"
    websites = "[.](com|net|org|io|gov)"

    def split_into_sentences(text):
        text = " " + text + "  "
        # R = dec.findall(text)
        # S = [j[0]+'<decimal>'+j[1] for j in [dig.findall(i) for i in R ]]
        # for i in range(len(R)):
        #  text = text.replace(R[i],S[i])
        text = text.replace("\n", " ")
        text = text.replace("Œ", "<stop>")
        text = re.sub(prefixes, "\\1<prd>", text)
        text = re.sub(websites, "<prd>\\1", text)
        if "Ph.D" in text: text = text.replace("Ph.D.", "Ph<prd>D<prd>")
        text = re.sub("\s" + alphabets + "[.] ", " \\1<prd> ", text)
        text = re.sub(acronyms + " " + starters, "\\1<stop> \\2", text)
        text = re.sub(alphabets + "[.]" + alphabets + "[.]" + alphabets + "[.]", "\\1<prd>\\2<prd>\\3<prd>", text)
        text = re.sub(alphabets + "[.]" + alphabets + "[.]", "\\1<prd>\\2<prd>", text)
        text = re.sub(" " + suffixes + "[.] " + starters, " \\1<stop> \\2", text)
        text = re.sub(" " + suffixes + "[.]", " \\1<prd>", text)
        text = re.sub(" " + alphabets + "[.]", " \\1<prd>", text)
        if "”" in text: text = text.replace(".”", "”.")
        if "\"" in text: text = text.replace(".\"", "\".")
        if "!" in text: text = text.replace("!\"", "\"!")
        if "?" in text: text = text.replace("?\"", "\"?")
        text = text.replace(".", ".<stop>")
        text = text.replace("?", "?<stop>")
        text = text.replace("!", "!<stop>")
        text = text.replace("<prd>", ".")
        sentences = text.split("<stop>")
        # sentences = sentences[:-1]
        sentences = [s.strip() for s in sentences]
        return sentences

    def readtxt(filename):
        doc = docx.Document(filename)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        return '\n'.join(fullText)

    # _______________________________________________
    #     WordSearch / PATTERN SEARCH FROM SENTENCES
    def search(word, sentences):
        return [i for i in sentences if re.search(r'\b%s\b' % word, i)]

    # _______________________________________________
    #
    # This function is taking Key and dataframe and returning all the rows and columns with matched Keywords
    def xsearch(regex: str, df, case=False):
        """Search all the text columns of `df`, return rows with any matches."""
        textlikes = df.select_dtypes(include=[object, "string"])
        return df[
            textlikes.apply(
                lambda column: column.str.contains(regex, regex=True, case=case, na=False)
            ).any(axis=1)]

    # _______________________________________________
    # This function is searching the Key inside the filtered list
    def indexHaveSubstring(lst, substring):
        ls = list()
        for i in range(len(lst)):
            subs = re.compile(Ke).search(str(lst[i]))
            if subs != None:
                ls.append(i)
        return ls

    p1 = File_pdf.objects.values('notesfile1')
    p2 = pd.DataFrame(p1, columns=['notesfile1'])
    # print(len(p2.iloc[:]['notesfile']))
    F_Count = len(p2.iloc[:]['notesfile1'])
    # print(F_Count)
    # print(p2)
    # KW = KW1.at[:,'keywords']
    # p3 = p2.iloc[-1]['notesfile']
    # p3 = p2
    print(p2)
    for p3 in p2.iloc[:]['notesfile1']:
        # print(p3)
        p = "C://Users/sumithra.r/Downloads/deployment/media/" + p3
        # BASE_DIR1 = os.path.join(BASE_DIR, 'media')
        # p = BASE_DIR1+p3
        print(p)
    # p = r'C:\\Users\sumithra.r\OneDrive - Concentrix Corporation\Handwritten Text\SAH-Outpatient-Therapy-Patient-Form-Jan-2018.pdf'
    # Text = textract.process(p)
    import zlib
    pdfFileObj = open(p, 'r+b')  # rb
    p1 = p
    pdfReader = PyPDF4.PdfFileReader(pdfFileObj,
                                     strict=False)  ###   STRICT + FALSE ADDED TO REMOVE THE ERROR CHECK FOR ACCURACY
    # pdfReader.save('test.docx')
    if not pdfReader.isEncrypted:
        tXt = "";
        PGtXt = ""
        for PG in range(pdfReader.numPages):
            PGNo = PG
            pageObj = pdfReader.getPage(PG)
            PGtXt = pageObj.extractText()
            PGtXt = PGtXt.replace(' !"#$%&\'()*+,-./012-345627\n3,&"8%))9"#:;3<%&$9=%)35\'&%%>%#?\n3,:\'%\n3@3',
                                  '').replace('!%3', '').replace('!&3', '').replace('!$3', '').replace('!#3',
                                                                                                       '').replace(
                '!"3', '').replace('!!3', '').replace('!*3', '').replace('!$3',
                                                                         '').replace('!*3', '').replace('!3',
                                                                                                        '').replace(
                '#3', '').replace('$3', '').replace('%3', '').replace("'3", '').replace('(3', '').replace(')3',
                                                                                                          '').replace(
                '&3', '')
            tXt = tXt + '\n-\r-\n' + PGtXt
            # PGtXt.save('test.docx')
        a_list = tXt
        excel_filename2 = "pdf2word" + str(datetime.datetime.today().date()) + ".docx"
        textfile = open(excel_filename2, "w", encoding="utf-8")
        # for element in a_list:
        result = textfile.write(a_list)
        print(textfile)
        # textfile.close()

        # print(excel_filename2)
        # textfile.write(BASE_DIR + '/media/' + excel_filename2)

        # df_filepath_out = '/media/' + textfile
    return render(request, 'CNXProtect_Utl/pdfword.html', {'df_filepath_out': result})

from django.shortcuts import render
from django.http import HttpResponse
from django.contrib.auth.models import User, auth
from django.contrib import messages
from django.shortcuts import render, redirect

# -*- coding: utf-8 -*-
"""
Created on Wed Aug 11 10:31:59 2021

@author: RNALAB
"""
# IMPORT PACKAGES

# import win32com.client

import ast
import os
import re
import numpy as np
import numpy
import docx
# import numpy
import PyPDF4
# import PyPDF2
import easygui
# import textract
import pandas as pd
from itertools import compress
from zipfile import ZipFile
# import win32com.client
import nltk
from nltk.stem import WordNetLemmatizer
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
# from string import punctuation
import spacy
import string
from spacy.lang.en.stop_words import STOP_WORDS
import time

start_time = time.time()

'''     INITIALIZATIONS     '''
dec = re.compile('\d+\.\d+')
dig = re.compile('\d+')
alphabets = "([A-Za-z])"
prefixes = "(Mr|St|Mrs|Ms|Dr)[.]"
suffixes = "(Inc|Ltd|Jr|Sr|Co)"
starters = "(Mr|Mrs|Ms|Dr|He\s|She\s|It\s|They\s|Their\s|Our\s|We\s|But\s|However\s|That\s|This\s|Wherever)"
acronyms = "([A-Z][.][A-Z][.](?:[A-Z][.])?)"
websites = "[.](com|net|org|io|gov)"


def cl_files(request, context=None):
    global Statement1, Ph_number, emails
    import json
    # global df, PP6, q, Statement, Frq, No_of_word, No_of_Line

    def split_into_sentences(text):
        text = " " + text + "  "
        R = dec.findall(text)
        S = [j[0] + '<decimal>' + j[1] for j in [dig.findall(i) for i in R]]
        for i in range(len(R)):
            text = text.replace(R[i], S[i])
        text = text.replace("\n", " ")
        text = text.replace("Œ", "<stop>")
        text = re.sub(prefixes, "\\1<prd>", text)
        text = re.sub(websites, "<prd>\\1", text)
        if "Ph.D" in text: text = text.replace("Ph.D.", "Ph<prd>D<prd>")
        text = re.sub("\s" + alphabets + "[.] ", " \\1<prd> ", text)
        text = re.sub(acronyms + " " + starters, "\\1<stop> \\2", text)
        text = re.sub(alphabets + "[.]" + alphabets + "[.]" + alphabets + "[.]", "\\1<prd>\\2<prd>\\3<prd>", text)
        text = re.sub(alphabets + "[.]" + alphabets + "[.]", "\\1<prd>\\2<prd>", text)
        text = re.sub(" " + suffixes + "[.] " + starters, " \\1<stop> \\2", text)
        text = re.sub(" " + suffixes + "[.]", " \\1<prd>", text)
        text = re.sub(" " + alphabets + "[.]", " \\1<prd>", text)
        if "”" in text: text = text.replace(".”", "”.")
        if "\"" in text: text = text.replace(".\"", "\".")
        if "!" in text: text = text.replace("!\"", "\"!")
        if "?" in text: text = text.replace("?\"", "\"?")
        text = text.replace(".", ".<stop>")
        text = text.replace("?", "?<stop>")
        text = text.replace("!", "!<stop>")
        text = text.replace("<prd>", ".")
        sentences = text.split("<stop>")
        # sentences = sentences[:-1]
        sentences = [s.strip() for s in sentences]
        return sentences

    def readtxt(filename):
        doc = docx.Document(filename)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        return '\n'.join(fullText)

    # _______________________________________________
    #     WordSearch / PATTERN SEARCH FROM SENTENCES
    def search(word, sentences):
        return [i for i in sentences if re.search(r'\b%s\b' % word, i)]

    # _______________________________________________
    #
    # This function is taking Key and dataframe and returning all the rows and columns with matched Keywords
    def xsearch(regex: str, df, case=False):
        """Search all the text columns of `df`, return rows with any matches."""
        textlikes = df.select_dtypes(include=[object, "string"])
        return df[
            textlikes.apply(
                lambda column: column.str.contains(regex, regex=True, case=case, na=False)
            ).any(axis=1)]

    # _______________________________________________
    # This function is searching the Key inside the filtered list
    def indexHaveSubstring(lst, substring):
        ls = list()
        for i in range(len(lst)):
            subs = re.compile(Ke).search(str(lst[i]))
            if subs != None:
                ls.append(i)
        return ls

    # _______________________________________________
    #####     DATE SEARCH PATTERNS
    import calendar
    dat = re.compile(
        r'(?:\d{1,2}(?:(?:-|/)|(?:th|st|nd|rd)?\s))?(?:(?:(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)(?:(?:-|/)|(?:,|\.)?\s)?)?(?:\d{1,2}(?:(?:-|/)|(?:th|st|nd|rd)?\s))?)(?:\d{2,4})')
    dat = r'(?:\d{1,2}[-/th|st|nd|rd\s]*)?(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)|jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)[a-z\s,.]*(?:\d{1,2}[-/th|st|nd|rd)\s,]*)+(?:\d{3})+'
    dat = r'(?:\d{1,2}[-/th|st|nd|rd\s]*|[-/-])?(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)|jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)|\d{1,2}?)[a-z\s,.]*(?:\d{1,2}[-/th|st|nd|rd)\s,]*)+(?:\d{3})+'
    full_months = [month for month in calendar.month_name if month]
    short_months = [d[:3] for d in full_months]
    months = '|'.join(short_months + full_months)
    sep = r'[.,]?\s+'  # seperators
    day = r'\d+'
    year = r'\d+'
    day_or_year = r'\d+(?:\w+)?'
    KW2 = Dictionary.objects.values('keywords')
    KW1 = pd.DataFrame(KW2, columns=['keywords'])
    # KW = KW1.at[:,'keywords']
    Dictionary_Word = KW1
    print(Dictionary_Word)
    # print(KW1)
    flg = "lower"
    DF = pd.DataFrame(
        columns=["File_Name", "file_list", "Word", "Sentence", "Sentence_No", "Page_No", "Text_Summerization",
                 "Frequency_of_Keyword", "No_of_word_in_the_paragraph",
                 "No_of_Line_in_the_parapraph", "Few_words", "Email_ID", "Ph_number", "Tagged_word"])
    # print(DF)
    df1 = pd.DataFrame()
    FileTrack = pd.DataFrame(columns=["FileName", "Page", "Status", "CharCount"])
    # excel_filename2 = "Contract_File_output" + str(datetime.datetime.today().date()) + ".xlsx"
    # p = "C:\\Users\RNALAB\Concentrix Corporation\Contract Lists - Contract Library\Broadcom (CA Inc.)\CA Inc-CNX US(MSA)07-24-20exe.pdf"
    p1 = File.objects.values('notesfile')
    p2 = pd.DataFrame(p1, columns=['notesfile'])
    # print(len(p2.iloc[:]['notesfile']))
    F_Count = len(p2.iloc[:]['notesfile'])
    # print(F_Count)
    # print(p2)
    # KW = KW1.at[:,'keywords']
    # p3 = p2.iloc[-1]['notesfile']
    # p3 = p2
    print(p2)
    for p3 in p2.iloc[:]['notesfile']:
        # print(p3)
        pp = "C://Users/sumithra.r/Downloads/deployment/media/" + p3
        # BASE_DIR1 = os.path.join(BASE_DIR, 'media')
        # p = BASE_DIR1+p3
        print(pp)
        # File.objects.all().delete()
        # print(p)
        pdfFileObj = open(pp, "r+b")
        plen = 0
        p = pp
        p1 = p
        pdfReader = PyPDF4.PdfFileReader(pdfFileObj, strict=False)  ###   STRICT + FALSE ADDED TO REMOVE THE ERROR CHECK FOR ACCURACY
        if not pdfReader.isEncrypted:
            for k in range(0, Dictionary_Word.shape[1]):
                for j in range(0, Dictionary_Word.shape[0]):
                    Keyword_present = 1;
                    if str(Dictionary_Word.iloc[j, k]) != 'nan':
                        # for(m in splitted){
                        # print(names(Dictionary[k]))
                        #
                        q = Dictionary_Word.iloc[j, k]
                        print(q)
                        tXt = "";
                        PGtXt = ""
                        for PG in range(pdfReader.numPages):
                            PGNo = PG
                            print(PGNo)
                            pageObj = pdfReader.getPage(PG)
                            PGtXt = pageObj.extractText()
                            PGtXt = PGtXt.replace(
                                ' !"#$%&\'()*+,-./012-345627\n3,&"8%))9"#:;3<%&$9=%)35\'&%%>%#?\n3,:\'%\n3@3',
                                '').replace('!%3', '').replace('!&3', '').replace('!$3', '').replace('!#3', '').replace(
                                '!"3', '').replace('!!3', '').replace('!*3', '').replace('!$3',
                                                                                         '').replace('!*3', '').replace(
                                '!3', '').replace('#3', '').replace('$3', '').replace('%3', '').replace("'3",
                                                                                                        '').replace(
                                '(3', '').replace(')3', '').replace('&3', '')
                            tXt = tXt + '\n-\r-\n' + PGtXt
                            tXt = tXt.lower()
                            PGtXt = PGtXt.strip('\n');
                            PGtXt = PGtXt.replace('\n', '');
                            PGtXt = PGtXt.strip('\t');
                            PGtXt = PGtXt.replace('\t', ' ');
                            PGtXt = re.sub("\s+", ' ', PGtXt)
                            # print("before plen", plen);
                            plen = plen + len(PGtXt)
                            # print("after plen", plen);
                            PGt = split_into_sentences(PGtXt)
                            if (flg == "lower"):
                                patt = q.lower()
                                PGt = [x.lower() for x in PGt]
                            else:
                                patt = q;
                            ### tolower
                            word = patt
                            # print(q) #                    Test
                            if '+DATE' in q:
                                word = q.replace('+DATE', '')
                                word = word.lower()
                                PG1 = [[(word in x), PGt.index(x)] for x in PGt]
                                PG2 = list(compress(PGt, [item[0] for item in PG1]))
                                PG3 = [item[1] for item in PG1 if item[0] == True]
                                PP1 = [[bool(re.search(dat, x)), PG2.index(x)] for x in PG2]
                                PP2 = list(compress(PP1, [item[0] for item in PP1]))
                                PP3 = [item[1] for item in PP1 if item[0] == True]
                                PP4 = [PG2[item] for item in PP3]
                            elif '+' in patt:
                                word = word.split(sep='+')
                                dat = word[1]
                                PG1 = [[(word[0] in x), PGt.index(x)] for x in PGt]
                                PG2 = list(compress(PGt, [item[0] for item in PG1]))
                                PG3 = [item[1] for item in PG1 if item[0] == True]
                                PP1 = [[bool(dat in x), PG2.index(x)] for x in PG2]
                                PP2 = list(compress(PP1, [item[0] for item in PP1]))
                                PP3 = [item[1] for item in PP1 if item[0] == True]
                                PP4 = [PG2[item] for item in PP3]
                            else:
                                PG1 = [
                                    [((' ' + word + ' ' in x) or (word + ' ' in x) or (word + '.' in x)), PGt.index(x)]
                                    for x in PGt]
                                # PG2 = list(compress(PGt,[item[0] for item in PG1]))
                                PG2 = (compress(PGt, [item[0] for item in PG1]))
                                PG3 = [item[1] for item in PG1 if item[0] == True]
                                PP3 = PG3
                                import numpy as np
                                if PP3 != []:
                                    Pp3 = (np.array(PP3));
                                    print(len(Pp3));
                                    for lst in range(len(Pp3)):
                                        PP5 = Pp3[lst];
                                        # if (len(PGt[PP5]) < 250) and (PP5 < len(PGt) - 1):
                                        #     # Sentenses are:
                                        #     PP6 = PGt[PP5 - 1] + PGt[PP5] + PGt[PP5 + 1]
                                        # else:
                                        #     PP6 = PGt[PP5]
                                        # PP5 = Pp3[lst];
                                        if (len(PGt[PP5]) < 250) and (PP5 < len(PGt) - 1):
                                            # Sentenses are:
                                            PP6 = PGt[PP5 - 1] + PGt[PP5] + PGt[PP5 + 1]
                                        else:
                                            PP6 = PGt[PP5]
                                            # print(PGt)
                                        Keyword_present += Keyword_present;
                                        before_keyword, keyword, after_keyword = PP6.partition(q)
                                        l1 = []
                                        l1.append(before_keyword[0:20])
                                        l1.append(keyword)
                                        l1.append(after_keyword[0:70])
                                        # print("actual words", PP6)
                                        # print("Few words",l1)
                                        # Few_words = re.search(q, PP6)
                                        # print("Few words", Few_words)
                                        # Few_words = re.compile(q.lower()).search(PP6)
                                        # print('keyword:', q,'\n')
                                        # print("Sentence:", PP6,'\n')
                                        # text analysis starts here:PP
                                        Main_sentance = []
                                        sents = nltk.sent_tokenize(PP6.lower())
                                        No_of_word = len(sents[0])
                                        No_of_Line = int(No_of_word / 75)
                                        lemmatizer = WordNetLemmatizer()
                                        Ph_number = [];
                                        Tagged_word = [];
                                        for i in range(len(sents)):
                                            emails = re.findall(r"[a-z0-9\.\-+_]+@[a-z0-9\.\-+_]+\.[a-z]+", sents[i])
                                            # print(emails)
                                            if len(emails) != 0:
                                                Ph_number = re.findall(
                                                    r"((?:\+\d{2}[-\.\s]??|\d{4}[-\.\s]??)?(?:\d{3}[-\.\s]??\d{3}[-\.\s]??\d{4}|\(\d{3}\)\s*\d{3}[-\.\s]??\d{4}|\d{3}[-\.\s]??\d{4}))",
                                                    sents[i])
                                                words = nltk.word_tokenize(sents[i])
                                                POS = nltk.pos_tag(words)
                                                # print('POS tagging:', POS, '\n')
                                                Tagged_word = []
                                                for words, tag in POS:
                                                    if (tag == 'CD' or tag == 'NNS'):
                                                        Tagged_word.append(words)
                                                # print('Tagged_words:', Tagged_word, '\n')

                                            words = nltk.word_tokenize(sents[i])
                                            words = [lemmatizer.lemmatize(word) for word in words if
                                                     word not in set(stopwords.words('english'))]
                                            sents[i] = ' '.join(words)
                                            # punctuation = punctuation +'\n'
                                            for punctuation in string.punctuation:
                                                sents[i] = sents[i].replace(punctuation, '')
                                            # print('keyword:', q, '\n')
                                            # print('lemmatizer:', sents,'\n')
                                            # print('punctuation:', sents,'\n')
                                            wordss = nltk.word_tokenize(sents[i])
                                            POS = nltk.pos_tag(wordss)
                                            # print("\n")
                                            # print('POS tagging:', POS,'\n')
                                            Tagged_words = []
                                            for wordss, tag in POS:
                                                if (tag == 'NN' or tag == 'JJ'):
                                                    Tagged_words.append(wordss)

                                            # print('Tagged_words:', Tagged_words,'\n')
                                        # Frq = Tagged_words.count(q.lower())
                                        Sum_sentence = list()
                                        Key_list = sorted(q.split(), key=len)
                                        if len(Key_list) == 1:
                                            KEY = Key_list[0]
                                        else:
                                            KEY = Key_list[-1]

                                        for j in range(len(Tagged_words)):
                                            # subsent = re.compile(Key_list[-1].lower()).search(Tagged_words[j])
                                            subsent = re.compile(KEY.lower()).search(Tagged_words[j])
                                            # print(subsent)
                                            if subsent != None:
                                                Sum_sentence.append(j)
                                        Cm = Sum_sentence
                                        Column_size = Cm

                                        if Cm != "" and Cm != []:
                                            if (len(Cm) > 0):
                                                for i in range(len(Cm)):
                                                    SenTance = Tagged_words[Cm[i]]
                                                    SenTance1 = [];
                                                    SenTance2 = [];
                                                    Main_sentance1 = "";
                                                    Main_sentance = [];
                                                    if (Cm[i] < len(Tagged_words) - 2):
                                                        SenTance2_1 = Tagged_words[Cm[i] + 1];
                                                        SenTance2_2 = Tagged_words[Cm[i] + 2]
                                                        SenTance2.append(SenTance2_1)
                                                        SenTance2.append(SenTance2_2)
                                                    else:
                                                        SenTance2_3 = Tagged_words[Cm[i]]
                                                        SenTance2.append(SenTance2_3)
                                                    if (Cm[i] != 0):
                                                        SenTance1_1 = Tagged_words[Cm[i] - 1];
                                                        SenTance1_2 = Tagged_words[Cm[i] - 2];
                                                        SenTance1.append(SenTance1_2)
                                                        SenTance1.append(SenTance1_1)
                                                    else:
                                                        SenTance1_3 = Tagged_words[Cm[i]]
                                                        SenTance1.append(SenTance1_3)

                                                    # Main_sentance.append(SenTance3)
                                                    SENT_1 = SenTance1
                                                    # Main_sentance.append(SenTance)
                                                    SENT_2 = SenTance2
                                                    # Main_sentance =
                                                    Main_sentance.append(SENT_1)
                                                    # Main_sentance =
                                                    Main_sentance.append(SENT_2)
                                                    # Main_sentance.append(SenTance4)
                                                    del SenTance1, SenTance2
                                                    # print("\n")
                                        # Txt_Anlys = ' '.join(Main_sentance);
                                        Main_sentance1 = Main_sentance
                                        Txt_Anlys = numpy.asarray(Main_sentance1)
                                        Txt_Anlys1 = list(np.unique(Txt_Anlys))
                                        # print(numpy.asarray(Main_sentance1))
                                        Frq = Tagged_words.count(SenTance)
                                        if (".pdf" in p and len(tXt) < 150):
                                            ImagePDF = "Image PDF"
                                        else:
                                            ImagePDF = 'Proper PDF File'
                                        Statement = SenTance, "for", Txt_Anlys1;
                                        PP7 = PP6[1:80]
                                        # print(PP7)
                                        df = pd.Series(
                                            [p1, p3[6:], q, PP6, PP5, PG, Txt_Anlys, Frq, No_of_word,
                                             No_of_Line, l1, emails, Ph_number, Tagged_word],
                                            index=["File_Name", "file_list", "Word", "Sentence", "Sentence_No",
                                                   "Page_No", "Text_Summerization",
                                                   "Frequency_of_Keyword", "No_of_word_in_the_paragraph",
                                                   "No_of_Line_in_the_parapraph", "Few_words", "Email_ID", "Ph_number",
                                                   "Tagged_word"])
                                        df1 = df1.append(df, ignore_index=True)
                                        # print("Keyword:", q)
                                        # print(Statement)
                        if Keyword_present <= 1:
                            PP6 = "Keyword not present"
                            PP7 = "Keyword not present"
                            PP5 = "Nill"
                            l1 = "Keyword is not present"
                            PG = "Nill"
                            Txt_Anlys = "Keyword is not present for analysis"
                            emails = "Nill"
                            Ph_number = "Nill"
                            Statement = "Keyword is not present for analysis"
                            Frq = "Nill"
                            No_of_word = "Nil"
                            No_of_Line = "Nil"
                            Tagged_word = " Nil"

                            df = pd.Series(
                                [p1, p3[6:], q, PP6, PP5, PG, Txt_Anlys, Frq, No_of_word,
                                 No_of_Line, l1, emails, Ph_number, Tagged_word],
                                index=["File_Name", "file_list", "Word", "Sentence", "Sentence_No", "Page_No",
                                       "Text_Summerization",
                                       "Frequency_of_Keyword", "No_of_word_in_the_paragraph",
                                       "No_of_Line_in_the_parapraph", "Few_words", "Email_ID", "Ph_number",
                                       "Tagged_word"])
                            df1 = df1.append(df, ignore_index=True)
                        else:
                            Statement1 = Statement;
                            Keyword = q;
                            Txt_Anlys1 = Txt_Anlys;
    DF = DF.append(df1)
    # df = pd.DataFrame(df1, index='Sentence')
    DF1 = DF.file_list.to_string(index=False).rjust(10)
    # print(DF1.str.strip()) .format( left_aligned)
    DF2 = DF.Word.to_string(index=False).upper().rjust(10)
    DF3 = DF.Few_words.to_string(index=False)
    # print(DF3)
    DF4 = DF.Text_Summerization.to_string(index=False).rjust(10)

    df2 = pd.Series([p3, "File", "Read", plen], index=["FileName", "Page", "Status", "CharCount"])
    FileTrack = FileTrack.append(df2, ignore_index=True)
    # print(FileTrack)
    # excel_filename = "Contract_File_Search" + str(p1) + "_" + str(datetime.datetime.today().date()) + ".xlsx"
    File.objects.all().delete()
    Dictionary.objects.all().delete()
    excel_filename1 = "Contract_File_Tracker" + str(datetime.datetime.today().date()) + ".xlsx"
    FileTrack.to_excel(BASE_DIR + '/media/' + excel_filename1)
    df_filepath_trk = '/media/' + excel_filename1

    excel_filename2 = "Contract_File_output" + str(datetime.datetime.today().date()) + ".xlsx"
    DF.to_excel(BASE_DIR + '/media/' + excel_filename2)
    df_filepath_out = '/media/' + excel_filename2

    excel_filename3 = "Contract_File_Tracker" + str(datetime.datetime.today().date()) + ".xlsx"
    FileTrack.to_excel(BASE_DIR + '/media/' + excel_filename3)
    df_filepath = '/media/' + excel_filename3
    # DF_File =DF.to_excel( "C:\\Users\RNALAB\Concentrix Corporation\CNX_Protect\media\result.xlsx")

    # return render(request, 'CNXProtect_Utl/cl_files.html', print({'result': DF}))
    # return render(request, 'CNXProtect_Utl/cl_files.html', {'DF': DF})
    # return render(request, 'CNXProtect_Utl/login_page.html', {'DF': DF})
    return render(request, 'CNXProtect_Utl/Contract_Library1.html',
                  {'F_Count': F_Count, 'DF': DF, 'q': q, 'p3': p3, 'PP6': PP6, 'Statement': Statement, 'PP7': PP7,
                   'excel_filename1': excel_filename1, 'l1': l1, 'DF1': DF1, 'DF2': DF2, 'DF3': DF3, 'DF4': DF4,
                   'emails': emails, 'df_filepath_trk': df_filepath_trk, 'df_filepath_out': df_filepath_out})
    # del statement, p3, PP6

    # # return redirect( 'CNXProtect_Utl/Contract_Library.html',
    #               {'DF': DF, 'q': q, 'p1': p1, 'plen': plen, 'PP6': PP6, 'Statement': Statement, 'Frq': Frq,
    #                'No_of_word': No_of_word, 'No_of_Line': No_of_Line, 'excel_filename1': excel_filename1,
    #                'df_filepath_trk': df_filepath_trk, 'df_filepath_out': df_filepath_out})


import os
import re
import docx
# import numpy
import PyPDF4
# import PyPDF2
import easygui
# import textract
import pandas as pd
from itertools import compress
from zipfile import ZipFile
# import win32com.client
# print("If `flg` not specified as 'lower', search algorithm will be case sensitive\n")
flg = "upper"
# def FMEA_Files(request):
#     FileCount = 0
#     flg = "lower1"
#
#     '''  TEST DIRECTORIES BELOW   '''
#     # C:\Users\sumithra.r\OneDrive - Concentrix Corporation\Documents\FMEA
#     x = [os.path.join(r, file) for r, d, f in
#          os.walk(r"C:\Users\sumithra.r\Concentrix Corporation\Bhawna Agnani - FMEA_Copy") for file in f]
#     # x = [os.path.join(r,file) for r,d,f in os.walk(r"C:\\Users\sumithra.r\OneDrive - Concentrix Corporation\Documents\Contract Lists - Contract Library") for file in f]
#     # x = [os.path.join(r,file) for r,d,f in os.walk("C:/Users/joshua.john/Concentrix Corporation/Account Management Operations (AMO) - Documents") for file in f]
#     # x = [os.path.join(r,file) for r,d,f in os.walk("C:/Users/joshua.john/OneDrive - Concentrix Corporation/A A PROJECTS/CNX - ROT v3/PythonDir/test/TEST2") for file in f]
#     # print(x)
#     ''' Obtain list of files '''
#     files = [f for f in os.listdir('.') if os.path.isfile(f)]
#     files = filter(lambda f: f.endswith(('.pdf', '.PDF')), x)
#
#     '''   EXTRACTING ALL FILES IN DIRECTORY TO LISTS   '''
#     ypdf = [k for k in x if k.endswith('.pdf')]  # pdf Alone
#     yPDF = [k for k in x if k.endswith('.PDF')]  # PDF Alone
#     ydoc = [k for k in x if k.endswith('.doc')]  # doc Alone
#     yDOC = [k for k in x if k.endswith('.DOC')]  # DOC Alone
#     ymsg = [k for k in x if k.endswith('.msg')]  # msg Alone
#     yMSG = [k for k in x if k.endswith('.MSG')]  # MSG Alone
#     ydocx = [k for k in x if k.endswith('.docx')]  # docx Alone
#     yDOCX = [k for k in x if k.endswith('.DOCX')]  # DOCX Alone
#     y2 = ypdf + yPDF
#     yDo = yDOCX + ydocx
#     yDox = ydoc + yDOC
#     '''   CAP - how many file you would like to search for   '''
#     '''   Change CAP if you want to define a maximum limit   '''
#     # Assign value to CAP
#     # CAP = len(y2)
#     '''##################### DICTIONARY FILES #####################'''
#     # Choose dictionary
#     # F = easygui.fileopenbox("Dictionary File")
#     # Dictionary = pd.read_excel(F,sheet_name=0)
#     Dictionary = pd.read_excel(r"C:\\Users\RNALAB\Concentrix Corporation\Keywords\Keyword File Sep,03_2021.xlsx", sheet_name=2)
#     # Dictionary = pd.read_excel(r"C:\\Users\sumithra.r\Concentrix Corporation\CNX Protect\Keywords\Keyword File June 24_2021.xlsx",sheet_name=1)
#     # Dictionary = pd.read_excel(r"C:\\Users\sumithra.r\OneDrive - Concentrix Corporation\Documents\OneDrive\Keyword File.xlsx",sheet_name=0)
#     ################################
#     '''     INITIALIZATIONS     '''
#     dec = re.compile('\d+\.\d+')
#     dig = re.compile('\d+')
#     alphabets = "([A-Za-z])"
#     prefixes = "(Mr|St|Mrs|Ms|Dr)[.]"
#     suffixes = "(Inc|Ltd|Jr|Sr|Co)"
#     starters = "(Mr|Mrs|Ms|Dr|He\s|She\s|It\s|They\s|Their\s|Our\s|We\s|But\s|However\s|That\s|This\s|Wherever)"
#     acronyms = "([A-Z][.][A-Z][.](?:[A-Z][.])?)"
#     websites = "[.](com|net|org|io|gov)"


def add_keywords(request):
    return render(request, 'CNXProtect_Utl/add_keywords.html')


def add_keywords_FMEA(request):
    return render(request, 'CNXProtect_Utl/add_keywords_FMEA.html')


def keyword_list2(request):
    context = {'keyword_list': Dictionary.objects.all()}
    print(context)
    # print(keyword_list)
    return render(request, 'CNXProtect_Utl/add_keywords_FMEA.html', context)


def insert_keyword_item2(request: HttpRequest):
    Keyword = Dictionary(keywords=request.POST['keywords'])
    Keyword.save()
    Keys = Keyword.keywords
    print(Keys)
    # print(Dictionary.objects.all())
    return redirect('add_keywords_FMEA')


def delete_keyword_item2(request, keyword_id):
    keyword_to_delete = Dictionary.objects.get(id=keyword_id)
    keyword_to_delete.delete()
    return redirect('add_keywords_FMEA')


def FM_Real(request):

    def split_into_sentences(text):
        text = " " + text + "  "
        R = dec.findall(text)
        S = [j[0] + '<decimal>' + j[1] for j in [dig.findall(i) for i in R]]
        for i in range(len(R)):
            text = text.replace(R[i], S[i])
        text = text.replace("\n", " ")
        text = text.replace("Œ", "<stop>")
        text = re.sub(prefixes, "\\1<prd>", text)
        text = re.sub(websites, "<prd>\\1", text)
        if "Ph.D" in text: text = text.replace("Ph.D.", "Ph<prd>D<prd>")
        text = re.sub("\s" + alphabets + "[.] ", " \\1<prd> ", text)
        text = re.sub(acronyms + " " + starters, "\\1<stop> \\2", text)
        text = re.sub(alphabets + "[.]" + alphabets + "[.]" + alphabets + "[.]", "\\1<prd>\\2<prd>\\3<prd>", text)
        text = re.sub(alphabets + "[.]" + alphabets + "[.]", "\\1<prd>\\2<prd>", text)
        text = re.sub(" " + suffixes + "[.] " + starters, " \\1<stop> \\2", text)
        text = re.sub(" " + suffixes + "[.]", " \\1<prd>", text)
        text = re.sub(" " + alphabets + "[.]", " \\1<prd>", text)
        if "”" in text: text = text.replace(".”", "”.")
        if "\"" in text: text = text.replace(".\"", "\".")
        if "!" in text: text = text.replace("!\"", "\"!")
        if "?" in text: text = text.replace("?\"", "\"?")
        text = text.replace(".", ".<stop>")
        text = text.replace("?", "?<stop>")
        text = text.replace("!", "!<stop>")
        text = text.replace("<prd>", ".")
        sentences = text.split("<stop>")
        # sentences = sentences[:-1]
        sentences = [s.strip() for s in sentences]
        return sentences


    def readtxt(filename):
        doc = docx.Document(filename)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        return '\n'.join(fullText)


    # _______________________________________________
    #     WordSearch / PATTERN SEARCH FROM SENTENCES
    def search(word, sentences):
        return [i for i in sentences if re.search(r'\b%s\b' % word, i)]


    # _______________________________________________
    #
    # This function is taking Key and dataframe and returning all the rows and columns with matched Keywords
    def xsearch(regex: str, df, case=False):
        """Search all the text columns of `df`, return rows with any matches."""
        textlikes = df.select_dtypes(include=[object, "string"])
        return df[
            textlikes.apply(
                lambda column: column.str.contains(regex, regex=True, case=case, na=False)
            ).any(axis=1)]

        # _______________________________________________
        # This function is searching the Key inside the filtered list


    def indexHaveSubstring(lst, substring):
        ls = list()
        for i in range(len(lst)):
            subs = re.compile(Ke).search(str(lst[i]))
            if subs != None:
                ls.append(i)
        return ls


    # _______________________________________________
    # ####     DATE SEARCH PATTERNS
    import calendar

    dat = re.compile(
        r'(?:\d{1,2}(?:(?:-|/)|(?:th|st|nd|rd)?\s))?(?:(?:(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)(?:(?:-|/)|(?:,|\.)?\s)?)?(?:\d{1,2}(?:(?:-|/)|(?:th|st|nd|rd)?\s))?)(?:\d{2,4})')
    dat = r'(?:\d{1,2}[-/th|st|nd|rd\s]*)?(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)|jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)[a-z\s,.]*(?:\d{1,2}[-/th|st|nd|rd)\s,]*)+(?:\d{3})+'
    dat = r'(?:\d{1,2}[-/th|st|nd|rd\s]*|[-/-])?(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)|jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)|\d{1,2}?)[a-z\s,.]*(?:\d{1,2}[-/th|st|nd|rd)\s,]*)+(?:\d{3})+'
    full_months = [month for month in calendar.month_name if month]
    short_months = [d[:3] for d in full_months]
    months = '|'.join(short_months + full_months)
    sep = r'[.,]?\s+'  # seperators
    day = r'\d+'
    year = r'\d+'
    day_or_year = r'\d+(?:\w+)?'
    # dat = re.compile(rf'(?:{day}{sep})?(?:{months}){sep}{day_or_year}(?:{sep}{year})?')
    ################################
    TEST = list()
    # OldList = read_excel('RawData/RawData.xlsx')
    OldList = ["a.pdf", "b.pdf",
               "c.pdf",
               "C:/Users/joshua.john/Concentrix Corporation/Account Management Operations (AMO) - Documents\\Contracts for Legal Team\\0 CIS programs\\Catholic Financial Life\\Transaction Documents\\Work Orders\\WO017\\Catholic Financial Life-CNX US-CIS(Work Order 017)10-23-17exe.pdf"]
    c0 = "character(0)"
    OutPut = pd.DataFrame(columns=['Sentances', 'Keywords', 'TotalCount', 'FileName', 'ImagePDF', 'Word'])
    y2 = ypdf + yPDF  # Combine PDF files later once all the list of pdfs in the library
    ### OpTion = int(input("Enter your inputs in numeric format for the below options:\n1, Keyword Search - Dictionary\n2, Paragraph search\n3, Padding search\n4, Single Keyword Search\n5, FMEA\n\n"))
    ImagePDF = ""

    a = ypdf + yPDF
    DF = pd.DataFrame(
        columns=["FileName/DirName", "ImagePDF", "Word", "Theme/Topic", "Sentence", "SentenceNo", "TotalCount",
                 "PageNo"])
    FileTrack = pd.DataFrame(columns=["FileName", "Page", "Status", "CharCount"])
    df1 = pd.DataFrame()
    df2 = pd.DataFrame()
    print(len(a))
    # %%
    '''   CODE FOR SEARCHING   '''
    import time
    start_time = time.time()
    CAP = len(y2) + 1
    zzz = 0
    y2 = ypdf + yPDF
    # option = 5
    # option = 5
    # print("*** FMEA - Dictionary***")
    # Reading all types of files from folders
    # Dictionary = pd.read_excel(easygui.fileopenbox(),sheet_name=0)
    # fils  = [os.path.join(r,file) for r,d,f in os.walk(r"C:\Users\joyeeta.mallik\Concentrix Corporation\Bhawna Agnani - FMEA") for file in f]

    KW2 = Dictionary1.objects.values('keywords')
    print(KW2)
    KW1 = pd.DataFrame(KW2, columns=['keywords'])
    # KW = KW1.at[:,'keywords']
    Dictionary_Word = KW1
    print(Dictionary_Word)
    p1 = File_1.objects.values('notesfile_1')
    p2 = pd.DataFrame(p1, columns=['notesfile_1'])
    # print(len(p2.iloc[:]['notesfile']))
    F_Count = len(p2.iloc[:]['notesfile_1'])
    # print(F_Count)
    # print(p2)
    # KW = KW1.at[:,'keywords']
    # p3 = p2.iloc[-1]['notesfile']
    # p3 = p2
    print(p2)
    for p3 in p2.iloc[:]['notesfile_1']:
        # print(p3)
        pp = "C://Users/sumithra.r/Downloads/deployment/media/" + p3
        # BASE_DIR1 = os.path.join(BASE_DIR, 'media')
        # p = BASE_DIR1+p3
        print(pp)
        # File.objects.all().delete()
        # print(p)
        # pdfFileObj = open(pp, 'rb')
        plen = 0
        p = pp
        p1 = p
        x = pp
        File1 = pd.DataFrame()
        FileTrack1 = pd.DataFrame()
        # x = 'C://Users/RNALAB/Concentrix Corporation/Files_execute/FMEA/ABBOTT - Malaysia.xlsx'
        xl = pd.ExcelFile(x)
        for i in range(len(xl.sheet_names)):
            file = pd.read_excel(x, sheet_name=xl.sheet_names[i])
            sheetNam = xl.sheet_names[i]
            print(sheetNam)
            for k in range(0, Dictionary_Word.shape[1]):
                for j in range(0, Dictionary_Word.shape[0]):
                    if str(Dictionary_Word.iloc[j, k]) != 'nan':
                        Key = Dictionary_Word.iloc[j, k]
                        # print(Key)
                        if flg == "lower":
                            Ke = Key.lower()
                            Ke = Ke + ' | ' + Ke + ' |' + Ke + '$'
                        else:
                            Ke = Key.lower()
                            Ke = Ke + ' | ' + Ke + ' |' + Ke + '$|' + Key + ' | ' + Key + ' |' + Key + '$'
                        file2 = xsearch(Key, file)
                        file3 = file2.values.tolist()
                        Rw = ''
                        Cm = ''
                        for l in range(len(file3)):
                            Rw = file2.index[l]
                            Cm = indexHaveSubstring(file3[l], Key)
                            Rw = file2.index[l]
                            Cm = indexHaveSubstring(file3[l], Key)
                            if Cm != "" and Cm != []:
                                # print(Key,Rw,Cm,x,sheetNam)
                                if (len(Cm) > 1):
                                    for i in range(len(Cm)):
                                        SenTance = file.iloc[Rw, Cm[i]]
                                        # print(SenTance)
                                        df2 = pd.Series(
                                            [Key, Dictionary_Word.columns[k], Rw, Cm[i], x, sheetNam, SenTance],
                                            index=["Key", "Theme/Topic", "RowNum", "ColNum", "FileName",
                                                   "SheetName", "Sentance"])
                                        File1 = File1.append(df2, ignore_index=True)
                                        print(File1)
                                        df_read = pd.Series([x, "File", "read", "-"],
                                                            index=["FileName", "Page", "Status", "CharCount"])
                                        FileTrack1 = FileTrack1.append(df_read, ignore_index=True)

            # from pyxlsb import open_workbook
            # try:
            #     for x in Fxls:
            #         print(str(Fxls.index(x)) + " / " + str(len(Fxls)))
            #         FName = x
            #         with open_workbook(x) as wb:
            #             for i in range(len(wb.sheets)):
            #                 file = pd.read_excel(x, sheet_name=i, engine="pyxlsb")
            #                 sheetNam = wb.sheets[i];  # print(sheetNam)
            #                 # Sentences,Keywords,TotalCount,
            #                 for k in range(0, Dictionary_Word.shape[1]):
            #                     for j in range(0, Dictionary_Word.shape[0]):
            #                         if str(Dictionary_Word.iloc[j, k]) != 'nan':
            #                             Key = Dictionary_Word.iloc[j, k]
            #                             if flg == "lower":
            #                                 Ke = Key.lower()
            #                                 Ke = Ke + ' | ' + Ke + ' |' + Ke + '$'
            #                             else:
            #                                 Ke = Key.lower()
            #                                 Ke = Ke + ' | ' + Ke + ' |' + Ke + '$|' + Key + ' | ' + Key + ' |' + Key + '$'
            #                             file2 = xsearch(Key, file)
            #                             file3 = file2.values.tolist()
            #                             Rw = ''
            #                             Cm = ''
            #                             for l in range(len(file3)):
            #                                 Rw = file2.index[l]
            #                                 Cm = indexHaveSubstring(file3[l], Key)
            #                                 if Cm != "" and Cm != []:
            #                                     # print(Key,Rw,Cm,x,sheetNam)
            #                                     if (len(Cm) > 1):
            #                                         for i in range(len(Cm)):
            #                                             SenTance = file.iloc[Rw, Cm[i]]
            #                                             df2 = pd.Series(
            #                                                 [Key, Dictionary_Word.columns[k], Rw, Cm[i], x, sheetNam,
            #                                                  SenTance, ],
            #                                                 index=["Key", "Theme/Topic", "RowNum", "ColNum", "FileName",
            #                                                        "SheetName", "Sentance"])
            #                                             File1 = File1.append(df2, ignore_index=True)
            #                                             df_read = pd.Series([x, "File", "read"],
            #                                                                 index=["FileName", "Page", "Status"])
            #                                             FileTrack1 = FileTrack1.append(df_read, ignore_index=True)
            # except:
            #     df_unread = pd.Series([x, "File", "Cant' be read"], index=["FileName", "Page", "Status"])
            #     FileTrack1 = FileTrack1.append(df_unread, ignore_index=True)
            #     print("File can't be read")
    # File_1.objects.all().delete()
    # Dictionary1.objects.all().delete()
    DF = DF.append(File1)

    FMDF1 = DF.Key.to_string(index=False)
    # print(DF1.str.strip()) .format( left_aligned)
    FMDF2 = DF.Sentance.to_string(index=False)
    FMDF3 = DF.FileName.to_string(index=False).rjust(10)
    # print(DF3)
    # File1.drop(labels=None, index=["Key", "Sentance", "FileName"])
    # DF4 = DF.Text_Summerization.to_string(index=False).rjust(10)
    excel_filename_2_1 = "FMEA_output" + str(datetime.datetime.today().date()) + ".xlsx"
    File1.to_excel(BASE_DIR + '/media/' + excel_filename_2_1)
    df_filepath_out1_1 = '/media/' + excel_filename_2_1

    return render(request, 'CNXProtect_Utl/FMEA_Real.html', {'df_filepath_out1_1': df_filepath_out1_1, 'FMDF1': FMDF1, 'FMDF2': FMDF2, 'FMDF3': FMDF3,})


def download(request, path):
    file_path = os.path.join(settings.MEDIA_ROOT, path)
    print(file_path)
    if os.path.exists(file_path):
        with open(file_path, 'rb') as fh:
            response = HttpResonse(fh.read(), content_type="application/notefile")
            response['Content-Disposition'] = 'inline;filename=' + os.path.basename(file_path)
            return response
    raise Http404


# IMPORT PACKAGES
import os
import re
import docx
# import numpy
import PyPDF4
# import PyPDF2
import easygui
# import textract
import pandas as pd
from itertools import compress
from zipfile import ZipFile
# import win32com.client

print("If `flg` not specified as 'lower', search algorithm will be case sensitive\n")
# flg = "lower1"
# # flg = "upper"
# FileCount = 0
'''  TEST DIRECTORIES BELOW   '''
x = [os.path.join(r, file) for r, d, f in os.walk(r"C:\Users\RNALAB\Concentrix Corporation\Contract_File") for file in
     f]
''' Obtain list of files '''
files = [f for f in os.listdir('.') if os.path.isfile(f)]
files = filter(lambda f: f.endswith(('.pdf', '.PDF')), x)
# File_Name = f
'''   EXTRACTING ALL FILES IN DIRECTORY TO LISTS   '''
ypdf = [k for k in x if k.endswith('.pdf')]  # pdf Alone
yPDF = [k for k in x if k.endswith('.PDF')]  # PDF Alone
ydoc = [k for k in x if k.endswith('.doc')]  # doc Alone
yDOC = [k for k in x if k.endswith('.DOC')]  # DOC Alone
ymsg = [k for k in x if k.endswith('.msg')]  # msg Alone
yMSG = [k for k in x if k.endswith('.MSG')]  # MSG Alone
ydocx = [k for k in x if k.endswith('.docx')]  # docx Alone
yDOCX = [k for k in x if k.endswith('.DOCX')]  # DOCX Alone
y2 = ypdf + yPDF
yDo = yDOCX + ydocx
yDox = ydoc + yDOC
Dictionary_Word = pd.read_excel(r"C:\Users\sumithra.r\Downloads\Docs files\Keyword_File_CL.xlsx", sheet_name=0)
dec = re.compile('\d+\.\d+')
dig = re.compile('\d+')
alphabets = "([A-Za-z])"
prefixes = "(Mr|St|Mrs|Ms|Dr)[.]"
suffixes = "(Inc|Ltd|Jr|Sr|Co)"
starters = "(Mr|Mrs|Ms|Dr|He\s|She\s|It\s|They\s|Their\s|Our\s|We\s|But\s|However\s|That\s|This\s|Wherever)"
acronyms = "([A-Z][.][A-Z][.](?:[A-Z][.])?)"
websites = "[.](com|net|org|io|gov)"


def split_into_sentences(text):
    text = " " + text + "  "
    R = dec.findall(text)
    S = [j[0] + '<decimal>' + j[1] for j in [dig.findall(i) for i in R]]
    for i in range(len(R)):
        text = text.replace(R[i], S[i])
    text = text.replace("\n", " ")
    text = text.replace("Œ", "<stop>")
    text = re.sub(prefixes, "\\1<prd>", text)
    text = re.sub(websites, "<prd>\\1", text)
    if "Ph.D" in text: text = text.replace("Ph.D.", "Ph<prd>D<prd>")
    text = re.sub("\s" + alphabets + "[.] ", " \\1<prd> ", text)
    text = re.sub(acronyms + " " + starters, "\\1<stop> \\2", text)
    text = re.sub(alphabets + "[.]" + alphabets + "[.]" + alphabets + "[.]", "\\1<prd>\\2<prd>\\3<prd>", text)
    text = re.sub(alphabets + "[.]" + alphabets + "[.]", "\\1<prd>\\2<prd>", text)
    text = re.sub(" " + suffixes + "[.] " + starters, " \\1<stop> \\2", text)
    text = re.sub(" " + suffixes + "[.]", " \\1<prd>", text)
    text = re.sub(" " + alphabets + "[.]", " \\1<prd>", text)
    if "”" in text: text = text.replace(".”", "”.")
    if "\"" in text: text = text.replace(".\"", "\".")
    if "!" in text: text = text.replace("!\"", "\"!")
    if "?" in text: text = text.replace("?\"", "\"?")
    text = text.replace(".", ".<stop>")
    text = text.replace("?", "?<stop>")
    text = text.replace("!", "!<stop>")
    text = text.replace("<prd>", ".")
    sentences = text.split("<stop>")
    # sentences = sentences[:-1]
    sentences = [s.strip() for s in sentences]
    return sentences


def readtxt(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

    # _______________________________________________
    #     WordSearch / PATTERN SEARCH FROM SENTENCES


def search(word, sentences):
    return [i for i in sentences if re.search(r'\b%s\b' % word, i)]

    # _______________________________________________
    #
    # This function is taking Key and dataframe and returning all the rows and columns with matched Keywords


def xsearch(regex: str, df, case=False):
    """Search all the text columns of `df`, return rows with any matches."""
    textlikes = df.select_dtypes(include=[object, "string"])
    return df[
        textlikes.apply(
            lambda column: column.str.contains(regex, regex=True, case=case, na=False)
        ).any(axis=1)]

    # _______________________________________________
    # This function is searching the Key inside the filtered list

def indexHaveSubstring(lst, substring):
    ls = list()
    for i in range(len(lst)):
        subs = re.compile(Ke).search(str(lst[i]))
        if subs != None:
            ls.append(i)
    return ls

    # _______________________________________________
    #####     DATE SEARCH PATTERNS
import calendar

dat = re.compile(
    r'(?:\d{1,2}(?:(?:-|/)|(?:th|st|nd|rd)?\s))?(?:(?:(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)(?:(?:-|/)|(?:,|\.)?\s)?)?(?:\d{1,2}(?:(?:-|/)|(?:th|st|nd|rd)?\s))?)(?:\d{2,4})')
dat = r'(?:\d{1,2}[-/th|st|nd|rd\s]*)?(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)|jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)[a-z\s,.]*(?:\d{1,2}[-/th|st|nd|rd)\s,]*)+(?:\d{3})+'
dat = r'(?:\d{1,2}[-/th|st|nd|rd\s]*|[-/-])?(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)|jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)|\d{1,2}?)[a-z\s,.]*(?:\d{1,2}[-/th|st|nd|rd)\s,]*)+(?:\d{3})+'
full_months = [month for month in calendar.month_name if month]
short_months = [d[:3] for d in full_months]
months = '|'.join(short_months + full_months)
sep = r'[.,]?\s+'  # seperators
day = r'\d+'
year = r'\d+'
day_or_year = r'\d+(?:\w+)?'
# dat = re.compile(rf'(?:{day}{sep})?(?:{months}){sep}{day_or_year}(?:{sep}{year})?')
################################

# %%
'''   CODE FOR SEARCHING   '''
# def Contract_Path(request):
#
#     return render(request, 'CNXProtect_Utl/Contract_Library.html', {'path': path})
def Contract_library(request):
    TEST = list()
    # x = [os.path.join(r, file) for r, d, f in os.walk(r"C:\Users\RNALAB\Concentrix Corporation\Contract_File") for file
    #      in f]
    # path = [os.path.join(r1, file) for r1, d1 in os.walk(r"C:\Users\RNALAB\Concentrix Corporation\Contract_File") for file
    #      in r1]
    # OldList = read_excel('RawData/RawData.xlsx')
    OldList = ["a.pdf", "b.pdf",
               "c.pdf",
               "C:/Users/joshua.john/Concentrix Corporation/Account Management Operations (AMO) - Documents\\Contracts for Legal Team\\0 CIS programs\\Catholic Financial Life\\Transaction Documents\\Work Orders\\WO017\\Catholic Financial Life-CNX US-CIS(Work Order 017)10-23-17exe.pdf"]
    c0 = "character(0)"
    OutPut = pd.DataFrame(columns=['Sentances', 'Keywords', 'TotalCount', 'FileName', 'ImagePDF', 'Word'])
    y2 = ypdf + yPDF  # Combine PDF files later once all the list of pdfs in the library
    # OpTion = int(input("Enter your inputs in numeric format for the below options:\n1, Keyword Search - Dictionary\n2, Paragraph search\n3, Padding search\n4, Single Keyword Search\n5, FMEA\n\n"))
    ImagePDF = ""
    # path = r"C:\Users\RNALAB\Concentrix Corporation\Contract_File"
    # print("file path:", path)
    a = ypdf + yPDF
    DF = pd.DataFrame(
        columns=["FileName/DirName", "ImagePDF", "Word", "Theme/Topic", "Sentence", "SentenceNo", "TotalCount",
                 "PageNo"])
    FileTrack = pd.DataFrame(columns=["FileName", "Page", "Status", "CharCount"])
    df1 = pd.DataFrame()
    df2 = pd.DataFrame()
    print(len(a))

    import time
    start_time = time.time()
    CAP = len(y2) + 1
    zzz = 0
    y2 = ypdf + yPDF
    # if (OpTion==1):
    print("*** Keyword Search - Using Dictionary ***")
    ### First n files
    # for i in range(0,((len(ypdf))-CAP)):
    # print(i)
    g = a[zzz:CAP]
    y2 = y2[zzz:CAP]
    ImagePDF = ""
    flg = "lower1"
    FileCount = 0
    if not (g in OldList):
        # print(a[i])
        # print(a[i])
        # FileRead[1,i] = append(FileRead,a[i])
        combined = ""
        # File = read_document(file = p,ocr = FALSE)##
        ###     Go Through List of All PDF files
        for p in y2:
            plen = 0
            if p != "C:/Users/joshua.john/Concentrix Corporation/Account Management Operations (AMO) - Documents\\Contracts for Legal Team\\0 CIS programs\\Catholic Financial Life\\Transaction Documents\\Work Orders\\WO017\\Catholic Financial Life-CNX US-CIS(Work Order 017)10-23-17exe.pdf":
                try:
                    FileCount = FileCount + 1
                    print(FileCount)
                    try:
                        pdfFileObj = open(p, 'r+b')  # rb
                        p1 = p
                    except:
                        df2 = pd.Series([p, "File", "Access Denied", plen],
                                        index=["FileName", "Page", "Status", "CharCount"])
                        FileTrack = FileTrack.append(df2, ignore_index=True)
                        print("File added as Access Denied")
                        pdfFileObj = open("C:/Users/joshua.john/Desktop/ROT Out/Blank/Blank.pdf", 'rb')
                        p1 = p + "BLANK FILE"
                    pdfReader = PyPDF4.PdfFileReader(pdfFileObj,
                                                     strict=False)  ###   STRICT + FALSE ADDED TO REMOVE THE ERROR CHECK FOR ACCURACY
                    if not pdfReader.isEncrypted:
                        tXt = "";
                        PGtXt = ""
                        for PG in range(pdfReader.numPages):
                            PGNo = PG
                            pageObj = pdfReader.getPage(PG)
                            try:
                                PGtXt = pageObj.extractText()
                            except:
                                PGtXt = ""
                                df2 = pd.Series([p1, PG + 1, "Unreadable", plen],
                                                index=["FileName", "Page", "Status", "CharCount"])
                                FileTrack = FileTrack.append(df2, ignore_index=True)
                                print("File added to unreadable list")
                            PGtXt = PGtXt.replace(
                                ' !"#$%&\'()*+,-./012-345627\n3,&"8%))9"#:;3<%&$9=%)35\'&%%>%#?\n3,:\'%\n3@3',
                                '').replace('!%3', '').replace('!&3', '').replace('!$3', '').replace('!#3', '').replace(
                                '!"3', '').replace('!!3', '').replace('!*3', '').replace('!$3',
                                                                                         '').replace('!*3', '').replace(
                                '!3', '').replace('#3', '').replace('$3', '').replace('%3', '').replace("'3",
                                                                                                        '').replace(
                                '(3', '').replace(')3', '').replace('&3', '')
                            tXt = tXt + '\n-\r-\n' + PGtXt

                            for k in range(0, Dictionary_Word.shape[1]):
                                for j in range(0, Dictionary_Word.shape[0]):
                                    if str(Dictionary_Word.iloc[j, k]) != 'nan':
                                        # for(m in splitted){
                                        # print(names(Dictionary[k]))
                                        #
                                        q = Dictionary_Word.iloc[j, k]
                                        # print(q)
                                        PGtXt = PGtXt.strip('\n');
                                        PGtXt = PGtXt.replace('\n', '');
                                        PGtXt = PGtXt.strip('\t');
                                        PGtXt = PGtXt.replace('\t', ' ');
                                        PGtXt = re.sub("\s+", ' ', PGtXt)
                                        plen = plen + len(PGtXt)
                                        PGt = split_into_sentences(PGtXt)
                                        if (flg == "lower"):
                                            patt = q.lower()
                                            PGt = [x.lower() for x in PGt]
                                        else:
                                            patt = q;
                                        ### tolower
                                        word = patt
                                        # print(q) #                    Test
                                        if '+DATE' in q:
                                            word = q.replace('+DATE', '')
                                            word = word.lower()
                                            PG1 = [[(word in x), PGt.index(x)] for x in PGt]
                                            PG2 = list(compress(PGt, [item[0] for item in PG1]))
                                            PG3 = [item[1] for item in PG1 if item[0] == True]
                                            PP1 = [[bool(re.search(dat, x)), PG2.index(x)] for x in PG2]
                                            PP2 = list(compress(PP1, [item[0] for item in PP1]))
                                            PP3 = [item[1] for item in PP1 if item[0] == True]
                                            PP4 = [PG2[item] for item in PP3]
                                        elif '+' in patt:
                                            word = word.split(sep='+')
                                            dat = word[1]
                                            PG1 = [[(word[0] in x), PGt.index(x)] for x in PGt]
                                            PG2 = list(compress(PGt, [item[0] for item in PG1]))
                                            PG3 = [item[1] for item in PG1 if item[0] == True]
                                            PP1 = [[bool(dat in x), PG2.index(x)] for x in PG2]
                                            PP2 = list(compress(PP1, [item[0] for item in PP1]))
                                            PP3 = [item[1] for item in PP1 if item[0] == True]
                                            PP4 = [PG2[item] for item in PP3]
                                        else:
                                            PG1 = [[((' ' + word + ' ' in x) or (word + ' ' in x) or (word + '.' in x)),
                                                    PGt.index(x)] for x in PGt]
                                            # PG2 = list(compress(PGt,[item[0] for item in PG1]))
                                            PG2 = (compress(PGt, [item[0] for item in PG1]))
                                            PG3 = [item[1] for item in PG1 if item[0] == True]
                                            PP3 = PG3
                                            import numpy as np
                                            if (PP3 != []) and str(Dictionary_Word.iloc[j, k]) != 'nan':
                                                Pp3 = (np.array(PP3));
                                                for lst in range(len(Pp3)):
                                                    PP5 = Pp3[lst];
                                                    if (len(PGt[PP5]) < 190) and (PP5 < len(PGt) - 1):
                                                        # LPP5 = PGt[PP5]
                                                        PP6 = PGt[PP5 - 1] + PGt[PP5] + PGt[PP5 + 1]
                                                    else:
                                                        PP6 = PGt[PP5]
                                                    if (".pdf" in p and len(tXt) < 150):
                                                        ImagePDF = "Image PDF"
                                                    else:
                                                        ImagePDF = 'Proper PDF File'
                                                    df = pd.Series(
                                                        [p1, ImagePDF, q, Dictionary_Word.columns[k], PP6, PP5, PG + 1],
                                                        index=["FileName/DirName", "ImagePDF", "Word", "Theme/Topic",
                                                               "Sentence", "SentenceNo", "TotalCount", "PageNo"])
                                                    df1 = df1.append(df, ignore_index=True)
                                                    print("Keyword:", q)
                                                    print("Sentence:", PP6)
                            DF = DF.append(df1)
                            df1 = pd.DataFrame()
                        df2 = pd.Series([p1, "File", "Read", plen], index=["FileName", "Page", "Status", "CharCount"])
                        FileTrack = FileTrack.append(df2, ignore_index=True)
                        #   print("File added to readable list")
                        # create path
                    else:
                        if pdfReader.isEncrypted:
                            import pikepdf
                            with pikepdf.open(p) as pdf:
                                # del pdf.pages[-1]
                                pdf.save("decrypted.pdf")

                                pdfFileObj = open("decrypted.pdf", "r+b")
                                pdfReader = PyPDF4.PdfFileReader(pdfFileObj, strict=False)
                                tXt = "";
                                PGtXt = ""
                                for PG in range(pdfReader.numPages):
                                    PGNo = PG
                                    pageObj = pdfReader.getPage(PG)
                                    try:
                                        PGtXt = pageObj.extractText()
                                    except:
                                        PGtXt = ""
                                        df2 = pd.Series([p1, PG + 1, "Unreadable", plen],
                                                        index=["FileName", "Page", "Status", "CharCount"])
                                        FileTrack = FileTrack.append(df2, ignore_index=True)
                                        print("File added to exception list")
                                    PGtXt = PGtXt.replace(
                                        ' !"#$%&\'()*+,-./012-345627\n3,&"8%))9"#:;3<%&$9=%)35\'&%%>%#?\n3,:\'%\n3@3',
                                        '').replace('!%3', '').replace('!&3', '').replace('!$3', '').replace('!#3',
                                                                                                             '').replace(
                                        '!"3', '').replace('!!3', '').replace('!*3', '').replace('!$3',
                                                                                                 '').replace('!*3',
                                                                                                             '').replace(
                                        '!3', '').replace('#3', '').replace('$3', '').replace('%3', '').replace("'3",
                                                                                                                '').replace(
                                        '(3', '').replace(')3', '').replace('&3', '')
                                    tXt = tXt + '\n-\r-\n' + PGtXt

                                    for k in range(0, Dictionary_Word.shape[1]):
                                        for j in range(0, Dictionary_Word.shape[0]):
                                            if str(Dictionary_Word.iloc[j, k]) != 'nan':
                                                # for(m in splitted){
                                                # print(names(Dictionary[k]))
                                                #
                                                q = Dictionary_Word.iloc[j, k]
                                                # print(q)
                                                PGtXt = PGtXt.strip('\n');
                                                PGtXt = PGtXt.replace('\n', '');
                                                PGtXt = PGtXt.strip('\t');
                                                PGtXt = PGtXt.replace('\t', ' ');
                                                PGtXt = re.sub("\s+", ' ', PGtXt)
                                                plen = plen + len(PGtXt)
                                                PGt = split_into_sentences(PGtXt)
                                                if (flg == "lower"):
                                                    patt = q.lower()
                                                    PGt = [x.lower() for x in PGt]
                                                else:
                                                    patt = q;
                                                ### tolower
                                                word = patt
                                                # print(q) #                    Test
                                                if '+DATE' in q:
                                                    word = q.replace('+DATE', '')
                                                    word = word.lower()
                                                    PG1 = [[(word in x), PGt.index(x)] for x in PGt]
                                                    PG2 = list(compress(PGt, [item[0] for item in PG1]))
                                                    PG3 = [item[1] for item in PG1 if item[0] == True]
                                                    PP1 = [[bool(re.search(dat, x)), PG2.index(x)] for x in PG2]
                                                    PP2 = list(compress(PP1, [item[0] for item in PP1]))
                                                    PP3 = [item[1] for item in PP1 if item[0] == True]
                                                    PP4 = [PG2[item] for item in PP3]
                                                elif '+' in patt:
                                                    word = word.split(sep='+')
                                                    dat = word[1]
                                                    PG1 = [[(word[0] in x), PGt.index(x)] for x in PGt]
                                                    PG2 = list(compress(PGt, [item[0] for item in PG1]))
                                                    PG3 = [item[1] for item in PG1 if item[0] == True]
                                                    PP1 = [[bool(dat in x), PG2.index(x)] for x in PG2]
                                                    PP2 = list(compress(PP1, [item[0] for item in PP1]))
                                                    PP3 = [item[1] for item in PP1 if item[0] == True]
                                                    PP4 = [PG2[item] for item in PP3]
                                                else:
                                                    PG1 = [[((' ' + word + ' ' in x) or (word + ' ' in x) or (
                                                            word + '.' in x)), PGt.index(x)] for x in PGt]
                                                    # PG2 = list(compress(PGt,[item[0] for item in PG1]))
                                                    PG2 = (compress(PGt, [item[0] for item in PG1]))
                                                    PG3 = [item[1] for item in PG1 if item[0] == True]
                                                    PP3 = PG3
                                                    import numpy as np
                                                    if (PP3 != []) and str(Dictionary_Word.iloc[j, k]) != 'nan':
                                                        Pp3 = (np.array(PP3));
                                                        for lst in range(len(Pp3)):
                                                            PP5 = Pp3[lst];
                                                            if (len(PGt[PP5]) < 190) and (PP5 < len(PGt) - 1):
                                                                # LPP5 = PGt[PP5]
                                                                PP6 = PGt[PP5 - 1] + PGt[PP5] + PGt[PP5 + 1]
                                                            else:
                                                                PP6 = PGt[PP5]
                                                            if (".pdf" in p and len(tXt) < 150):
                                                                ImagePDF = "Image PDF"
                                                            else:
                                                                ImagePDF = 'Proper PDF File'
                                                            df = pd.Series(
                                                                [p1, ImagePDF, q, Dictionary_Word.columns[k], PP6, PP5,
                                                                 PG + 1], index=["FileName/DirName", "ImagePDF", "Word",
                                                                                 "Theme/Topic", "Sentence",
                                                                                 "SentenceNo", "TotalCount", "PageNo"])
                                                            df1 = df1.append(df, ignore_index=True)
                                                            print("Keyword:", q)
                                                            print("Sentence:", PP6)
                                # df1['FileName/DirName'] = p
                                # df1["ImagePDF"] = ImagePDF
                                # df1["Word"] = q
                                # df1["Theme/Topic"] = q
                                # df1["TotalCount"] = 0
                                # ["FileName/DirName","ImagePDF","Word","Theme/Topic","Sentence","PageNo","TotalCount"]
                                # combine sections for each word###
                                DF = DF.append(df1)
                                df1 = pd.DataFrame()
                        df2 = pd.Series([p1, "File", "Read", plen], index=["FileName", "Page", "Status", "CharCount"])
                        FileTrack = FileTrack.append(df2, ignore_index=True)

                        # pdf.close()
                        pdfFileObj.close()
                        os.remove("decrypted.pdf")
                        print("File added as decrypted file")
                        # create path
                except:
                    PGtXt = ""
                    df2 = pd.Series([p1, PG + 1, "ReadError", plen], index=["FileName", "Page", "Status", "CharCount"])
                    FileTrack = FileTrack.append(df2, ignore_index=True)
                    print("File added to ReadError list")
                    # os.path.join('C:\Users\sumithra.r\Concentrix Corporation\CNX Protect\Output Files\Readerror', 'p')
                    pdfFileObj.close()

    ###################################################################
    # DOCX FILES
    for p in yDo:
        try:
            DoTxt = readtxt(p)
            p1 = p
        except:
            df2 = pd.Series([p, "File", "Cant' be read", "-"], index=["FileName", "Page", "Status", "CharCount"])
            FileTrack = FileTrack.append(df2, ignore_index=True)
            print("File can't be read")
            DoTxt = ""
            p1 = p + "BLANK FILE ADDED"
        PGtXt = DoTxt
        PGtXt = PGtXt.replace(' !"#$%&\'()*+,-./012-345627\n3,&"8%))9"#:;3<%&$9=%)35\'&%%>%#?\n3,:\'%\n3@3',
                              '').replace('!%3', '').replace('!&3', '').replace('!$3', '').replace('!#3',
                                                                                                   '').replace(
            '!"3', '').replace('!!3', '').replace('!*3', '').replace('!$3', '').replace('!*3', '').replace('!3',
                                                                                                           '').replace(
            '#3', '').replace('$3', '').replace('%3', '').replace("'3", '').replace('(3', '').replace(')3',
                                                                                                      '').replace(
            '&3', '')
        tXt = tXt + '\n-\r-\n' + PGtXt
        for k in range(0, Dictionary_Word.shape[1]):
            for j in range(0, Dictionary_Word.shape[0]):
                if str(Dictionary_Word.iloc[j, k]) != 'nan':
                    # for(m in splitted){
                    # print(names(Dictionary[k]))
                    #
                    q = Dictionary_Word.iloc[j, k]
                    # print(q)
                    PGtXt = PGtXt.strip('\n');
                    PGtXt = PGtXt.replace('\n', '');
                    PGtXt = PGtXt.strip('\t');
                    PGtXt = PGtXt.replace('\t', ' ');
                    PGtXt = re.sub("\s+", ' ', PGtXt)
                    plen = plen + len(PGtXt)
                    PGt = split_into_sentences(PGtXt)
                    if (flg == "lower"):
                        patt = q.lower()
                        PGt = [x.lower() for x in PGt]
                    else:
                        patt = q;
                    ### tolower
                    word = patt
                    # print(q) #                    Test
                    if '+DATE' in q:
                        word = q.replace('+DATE', '')
                        word = word.lower()
                        PG1 = [[(word in x), PGt.index(x)] for x in PGt]
                        PG2 = list(compress(PGt, [item[0] for item in PG1]))
                        PG3 = [item[1] for item in PG1 if item[0] == True]
                        PP1 = [[bool(re.search(dat, x)), PG2.index(x)] for x in PG2]
                        PP2 = list(compress(PP1, [item[0] for item in PP1]))
                        PP3 = [item[1] for item in PP1 if item[0] == True]
                        PP4 = [PG2[item] for item in PP3]
                    elif '+' in patt:
                        word = word.split(sep='+')
                        dat = word[1]
                        PG1 = [[(word[0] in x), PGt.index(x)] for x in PGt]
                        PG2 = list(compress(PGt, [item[0] for item in PG1]))
                        PG3 = [item[1] for item in PG1 if item[0] == True]
                        PP1 = [[bool(dat in x), PG2.index(x)] for x in PG2]
                        PP2 = list(compress(PP1, [item[0] for item in PP1]))
                        PP3 = [item[1] for item in PP1 if item[0] == True]
                        PP4 = [PG2[item] for item in PP3]
                    else:
                        PG1 = [[((' ' + word + ' ' in x) or (word + ' ' in x) or (word + '.' in x)), PGt.index(x)]
                               for x in PGt]
                        # PG2 = list(compress(PGt,[item[0] for item in PG1]))
                        PG2 = (compress(PGt, [item[0] for item in PG1]))
                        PG3 = [item[1] for item in PG1 if item[0] == True]
                        PP3 = PG3
                        import numpy as np
                        if (PP3 != []) and str(Dictionary_Word.iloc[j, k]) != 'nan':
                            Pp3 = (np.array(PP3));
                            for lst in range(len(Pp3)):
                                PP5 = Pp3[lst];
                                if (len(PGt[PP5]) < 190) and (PP5 < len(PGt) - 1):
                                    # LPP5 = PGt[PP5]
                                    PP6 = PGt[PP5 - 1] + PGt[PP5] + PGt[PP5 + 1]
                                else:
                                    PP6 = PGt[PP5]
                                if (".pdf" in p and len(tXt) < 150):
                                    ImagePDF = "Image PDF"
                                else:
                                    ImagePDF = 'Proper PDF File'
                                df = pd.Series([p1, ImagePDF, q, Dictionary_Word.columns[k], PP6, PP5, plen, PG + 1],
                                               index=["FileName/DirName", "ImagePDF", "Word", "Theme/Topic",
                                                      "Sentence", "SentenceNo", "TotalCount", "PageNo"])
                                df1 = df1.append(df, ignore_index=True)
                        # path = "C:\Users\RNALAB\Concentrix Corporation\Contract_File"\
                        # path = p1[1:40]
                        # import re
                        path = p1.split()[0:2]
                        print("path:", path)
                        print("Keyword:", q)
                        print("Sentence:", PP6)
        # df1['FileName/DirName'] = p
        # df1["ImagePDF"] = ImagePDF
        # df1["Word"] = q
        # df1["Theme/Topic"] = q
        # df1["TotalCount"] = 0
        # ["FileName/DirName","ImagePDF","Word","Theme/Topic","Sentence","PageNo","TotalCount"]

        # combine sections for each word###
        DF = DF.append(df1)
        # df1 = pd.DataFrame()
        ###################################################################
    # Word will change, for each word see for all combinations and apply conditions for multiple words
    # _____________________________________________________________________________
    DF = DF.reset_index();
    del (DF['index'])
    ccc = DF.groupby(['FileName/DirName', "Word"]).size().reset_index()
    for i in range(len(DF)):
        print(i)
        for j in range(len(ccc)):
            if (ccc['FileName/DirName'][j] == DF['FileName/DirName'][i] and ccc['Word'][j] == DF['Word'][i]):
                DF['TotalCount'][i] = ccc[0][j]

    FileTrack["Status1"] = ""
    for i in range(len(FileTrack)):
        print(i)
        if FileTrack["Status"][i] == "Read" and FileTrack["CharCount"][i] <= 150:
            FileTrack["Status1"][i] = 'Scanned File'
        else:
            FileTrack["Status1"][i] = FileTrack["Status"][i]
    excel_filename_1 = "Contract_Library_Tracker" + str(datetime.datetime.today().date()) + ".xlsx"
    FileTrack.to_excel(BASE_DIR + '/media/' + excel_filename_1)
    df_filepath_trk1 = '/media/' + excel_filename_1
    """   FILE WRITE OPERATION   """

    excel_filename_2 = "Contract_Library_output" + str(datetime.datetime.today().date()) + ".xlsx"
    DF.to_excel(BASE_DIR + '/media/' + excel_filename_2)
    df_filepath_out1 = '/media/' + excel_filename_2

    len(y2)
    excel_filename_3 = "Contract_Library_Execution" + str(datetime.datetime.today().date()) + ".xlsx"
    pd.DataFrame(y2).to_excel(BASE_DIR + '/media/' + excel_filename_3)
    df_filepath_Exe1 = '/media/' + excel_filename_3

    # path = "C:\Users\RNALAB\Concentrix Corporation\Contract_File"

    # DF_File =DF.to_excel( "C:\\Users\RNALAB\Concentrix Corporation\CNX_Protect\media\result.xlsx")
    return render(request, 'CNXProtect_Utl/Contract_Library1.html', {'excel_filename_1': excel_filename_1,
                                                                    'path': path,
                                                                    'df_filepath_trk1': df_filepath_trk1,
                                                                    'df_filepath_out1': df_filepath_out1,
                                                                    'df_filepath_Exe1': df_filepath_Exe1})




    # return render(request, 'CNXProtect_Utl/cl_files.html', print({'result': DF}))
    # return render(request, 'CNXProtect_Utl/cl_files.html', {'DF': DF})
    # return render(request, 'CNXProtect_Utl/login_page.html', {'path': p})
    # return render(request, 'CNXProtect_Utl/login_page.html', {'DF': DF, 'q': q, 'p1': p1, 'plen': plen, 'excel_filename_1': excel_filename_1,
    #                'df_filepath_trk1': df_filepath_trk1, 'df_filepath_out1': df_filepath_out1, 'df_filepath_Exe1': df_filepath_Exe1})
    #

    # return render(request, 'CNXProtect_Utl/cl_files.html', request.POST[DF])
    # return ()#print("--- %s minutes ---" % round((time.time() - start_time) / 60, 4))


def FMEA_Batch(request):
    return render(request, 'CNXProtect_Utl/FMEA.html')


def Main_page_template(request):
    return render(request, 'CNXProtect_Utl/Main_page_template.html')


def read_file_1(request):
    # if not request.user.is_authenticated:
    #     return redirect('login')
    # error = ""
    if request.method == "POST":
        # if request.POST:
        print("this is my first line")
        notesfile1 = request.FILES['notesfile1']
        p = notesfile1
        # p1 = file.notesfile
        a = File_pdf(notesfile1=notesfile1)
        a.save()
        p1 = File_pdf.objects.values('notesfile1')
        p2 = pd.DataFrame(p1, columns=['notesfile1'])
        F_Count = len(p2.iloc[:]['notesfile1'])
        print(F_Count)
        # print(p)        # p = File.objects.values('notesfile')
        # print(p)        # print("C:\Users\RNALAB\Concentrix Corporation\CNX_Protect\media\media" + p)
        messages.success(request,
                         'Files submitted successfully!!')  # return render(request, 'CNXProtect_Utl/keyword_list.html', p)
        return redirect('pdfword')
    else:
        messages.error(request, 'Files not submitted successfully!!')
        return redirect('indexing_data5')


def read_file(request):
    # if not request.user.is_authenticated:
    #     return redirect('login')
    # error = ""
    if request.method == "POST":
        # if request.POST:
        print("this is my first line")
        notesfile = request.FILES['notesfile']
        p = notesfile
        # p1 = file.notesfile
        a = File(notesfile=notesfile)
        a.save()
        p1 = File.objects.values('notesfile')
        p2 = pd.DataFrame(p1, columns=['notesfile'])
        F_Count = len(p2.iloc[:]['notesfile'])
        print(F_Count)
        # print(p)        # p = File.objects.values('notesfile')
        # print(p)        # print("C:\Users\RNALAB\Concentrix Corporation\CNX_Protect\media\media" + p)
        messages.success(request,
                         'Files submitted successfully!!')  # return render(request, 'CNXProtect_Utl/keyword_list.html', p)
        return redirect('keyword_list')
    else:
        messages.error(request, 'Files not submitted successfully!!')
        return redirect('indexing_data5')


def read_file1(request: HttpRequest):
    # if not request.user.is_authenticated:
    #     return redirect('login')
    # error = ""
    if request.method == "POST":
        # if request.POST:
        print("this is my first line")
        notesfile_1 = request.FILES['notesfile_1']
        p = notesfile_1
        # p1 = file.notesfile
        a = File_1(notesfile_1=notesfile_1)
        a.save()
        p1 = File_1.objects.values('notesfile_1')
        p2 = pd.DataFrame(p1, columns=['notesfile_1'])
        F_Count = len(p2.iloc[:]['notesfile_1'])
        print(F_Count)
        # print(p)        # p = File.objects.values('notesfile')
        # print(p)        # print("C:\Users\RNALAB\Concentrix Corporation\CNX_Protect\media\media" + p)
        messages.success(request,
                         'Files submitted successfully!!')  # return render(request, 'CNXProtect_Utl/keyword_list.html', p)
        return redirect('keyword_list1')
    else:
        messages.error(request, 'Files not submitted successfully!!')
        return redirect('indexing_data5')


def MSA_list(request):
    context = {'MSA_list': MSA.objects.all()}
    print(context)
    # print(keyword_list)
    return render(request, 'CNXProtect_Utl/Contract_Library1.html', context)


def insert_MSA_item(request: HttpRequest):
    Keyss = MSA(keywords=request.POST['account'])
    Keyss.save()
    Keys = Keyss.account
    print(Keys)
    # print(Dictionary.objects.all())
    return redirect('MSA_list')


def master_data(request):
    if request.method == "POST":
        notesfile_1 = request.FILES['notesfile_1']
        p = notesfile_1
        # p1 = file.notesfile
        a = Master_Data(notesfile_1=notesfile_1)
        a.save()
    p3 = Master_Data.objects.values('notesfile_1')
    # print(Master_Data.objects.filter(notesfile_1__startwith = '0' ))
    p4 = pd.DataFrame(p3, columns=['notesfile_1'])
    print(p4)
    F_Len = len(p4.iloc[:]['notesfile_1'])
    print(F_Len)
    F_Count = p4.iloc[F_Len - 1]['notesfile_1']
    # sumithra.r/Downloads/deployment/
    X_File = "C://Users/sumithra.r/Downloads/deployment/media/" + F_Count
    print(X_File)
    XL_File = pd.read_excel(X_File)
    MSA_Name = 'Acer'
    # request.GET['MSA_1']

    # print(pd.DataFrame(MSA_Name))
    Account = XL_File['MSA'].unique()
    DF11 = XL_File['MSA'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == 'Health Insurance')]
    DF12 = XL_File['Word'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == 'Health Insurance')]
    DF13 = XL_File['Sentence'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == 'Health Insurance')]
    DF14 = XL_File['Theme/Topic'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == 'Health Insurance')]
    DF15 = XL_File['File_Name/DirName'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == 'Health Insurance')]
    DF15df = pd.DataFrame(DF15)
    DF16 = XL_File['Date_of_creation'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == 'Health Insurance')]

    DF17 = XL_File.loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == 'Health Insurance')]
    excel_filename_1 = "Indexing_report_" + str(datetime.datetime.today().date()) + ".xlsx"
    DF17.to_excel(BASE_DIR + '/media/' + excel_filename_1)
    DF17_1 = '/media/' + excel_filename_1
    print(DF17)

    DF1_1 = DF11.to_string(index=False).rjust(10)
    DF1_2 = DF12.to_string(index=False).rjust(10)
    DF1_3 = DF13.to_string(index=False).rjust(10)
    DF1_4 = DF14.to_string(index=False).rjust(10)
    DF1_5 = DF15.to_string(index=False).rjust(10)
    DF1_6 = DF16.to_string(index=False).rjust(10)

    return render(request, 'CNXProtect_Utl/Indexing.html',
                  {'Account': Account, 'DF17_1': DF17_1, 'DF1_1': DF1_1, 'DF1_2': DF1_2, 'DF1_3': DF1_3, 'DF1_4': DF1_4,
                   'DF1_5': DF1_5, 'DF1_6': DF1_6})


def indexing1(request):
    # Keyword = Dictionary(keywords=request.POST['keywords'])
    # Keyword.save()
    # Keyword = Add_Keywords(Name=request.POST['Name'])
    # Keyword.save()
    # context = {'New_keyword': Add_Keywords.objects.all()}
    # print(context)
    # if request.method == "POST":
    #     notesfile_1 = request.FILES['notesfile_1']
    #     p = notesfile_1
    return render(request, "CNXProtect_Utl/indexing1.html")


def indexing_data2(request):
    if request.method == "POST":
        notesfile_11 = request.FILES['notesfile_11']
        p = notesfile_11
        # p1 = file.notesfile
        a = Master_Data1(notesfile_11=notesfile_11)
        a.save()
    p5 = Master_Data1.objects.values('notesfile_11')
    # print(Master_Data.objects.filter(notesfile_1__startwith = '0' ))
    p6 = pd.DataFrame(p5, columns=['notesfile_11'])
    print(p6)
    F_Len = len(p6.iloc[:]['notesfile_11'])
    print(F_Len)
    F_Count_1 = p6.iloc[F_Len - 1]['notesfile_11']
    F_Count1_1 = p6.iloc[F_Len - 2]['notesfile_11']
    X_File_1 = "C://Users/sumithra.r/Downloads/deployment/media/" + F_Count_1
    File_Tracker_1 = "C://sumithra.r/Downloads/deployment/media/" + F_Count1_1
    print("The output Master File is", X_File_1)
    print(X_File_1)
    XL_File_2 = pd.read_excel(X_File_1)
    # XL_File['Date_of_modified'] = pd.to_datetime(XL_File['Date_of_modified'])
    MSA_Name = request.GET.getlist('MSA_1')
    print("selected Theme:", MSA_Name)
    Theme_Name = request.GET.getlist('Theme')
    Geo_Name = request.GET.getlist('Geo')
    print("Selected Key:", Theme_Name)

    DF17 = XL_File_2[XL_File_2['MSA'].isin(MSA_Name) & XL_File_2['Key'].isin(Theme_Name) & XL_File_2['Geo'].isin(Geo_Name)]
    # DF17 = XL_File.loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]
    excel_filename_1 = "FMEA_Results_" + str(datetime.datetime.today().date()) + ".xlsx"
    DF17.to_excel(BASE_DIR + '/media/' + excel_filename_1)
    DF17_1 = '/media/' + excel_filename_1
    print(DF17)

    # print(pd.DataFrame(MSA_Name))
    Account = XL_File_2['Theme/Topic'].unique()
    # DF11 = XL_File['MSA'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Theme/Topic'] == mask)]
    # (df['birth_date'] > start_date) & (df['birth_date'] <= end_date)
    # DF12 = XL_File['Word'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]
    # DF13 = XL_File['Sentence'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]
    # DF14 = XL_File['Theme/Topic'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]
    # DF15 = XL_File['File_Name/DirName'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]
    # # DF15df = pd.DataFrame(DF15)
    # DF16 = XL_File['Date_of_modified'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]
    # print(DF13)
    # DF1_1 = DF11.to_string(index=False).rjust(10)
    # DF1_2 = DF12.to_string(index=False).rjust(10)
    # if DF13.empty:
    #     DF1_3 = 'No sentences have been extracted from the Account:', MSA_Name, 'and the Theme:', Theme_Name
    # else:
    #     DF1_3 = DF13.to_string(index=False).rjust(10)
    # DF1_4 = DF14.to_string(index=False).rjust(10)
    # DF1_5 = DF15.to_string(index=False).rjust(10)
    # DF1_6 = DF16.to_string(index=False).rjust(10)

    return render(request, 'CNXProtect_Utl/indexing2.html',
                  {'Account': Account, 'DF17_1_': DF17_1})
                  # {'Account': Account, 'DF17_1': DF17_1,
                  #      'DF1_1': DF1_1, 'DF1_2': DF1_2, 'DF1_3': DF1_3, 'DF1_4': DF1_4, 'DF1_5': DF1_5, 'DF1_6': DF1_6})


def indexing_data3(request):
    # Keyword = Dictionary(keywords=request.POST['keywords'])
    # Keyword.save()
    # Keyword = Add_Keywords(Name=request.POST['Name'])
    # Keyword.save()
    context = {'New_keyword': Add_Keywords.objects.all()}
    print(context)
    # if request.method == "POST":
    #     notesfile_1 = request.FILES['notesfile_1']
    #     p = notesfile_1
    return render(request, "CNXProtect_Utl/indexing3.html")


def indexing_data4(request):
    if request.method == "POST":
        notesfile_1 = request.FILES['notesfile_1']
        p = notesfile_1
        # p1 = file.notesfile
        a = Master_Data(notesfile_1=notesfile_1)
        a.save()
    p3 = Master_Data.objects.values('notesfile_1')
    # print(Master_Data.objects.filter(notesfile_1__startwith = '0' ))
    p4 = pd.DataFrame(p3, columns=['notesfile_1'])
    print(p4)
    F_Len = len(p4.iloc[:]['notesfile_1'])
    print(F_Len)
    F_Count = p4.iloc[F_Len - 1]['notesfile_1']
    X_File = "C://Users/sumithra.r/Downloads/deployment/media/" + F_Count
    # print(X_File)
    XL_File = pd.read_excel(X_File)
    # XL_File['Date_of_modified'] = pd.to_datetime(XL_File['Date_of_modified'])
    MSA_Name = request.GET.getlist('MSA_1')
    print(MSA_Name)
    Theme_Name = request.GET.getlist('Theme')
    # Time_in = request.GET['InputTime']
    # Time_out = request.GET['OutputTime']
    # mask = (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)
    # print(mask)
    # print(MSA_Name[0])
    # print(Theme_Name[0])
    # print(Time_in)
    # print(Time_out)

    DF17 = XL_File[XL_File['MSA'].isin(MSA_Name) & XL_File['Word'].isin(Theme_Name)]
    # DF17 = XL_File.loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]
    excel_filename_1 = "Indexing_report_" + str(datetime.datetime.today().date()) + ".xlsx"
    DF17.to_excel(BASE_DIR + '/media/' + excel_filename_1)
    DF17_1 = '/media/' + excel_filename_1
    print(DF17)

    # print(pd.DataFrame(MSA_Name))
    Account = XL_File['MSA'].unique()
    # DF11 = XL_File['MSA'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Theme/Topic'] == mask)]
    #                                                                                   # (df['birth_date'] > start_date) & (df['birth_date'] <= end_date)
    # DF12 = XL_File['Word'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]
    # DF13 = XL_File['Sentence'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]
    # DF14 = XL_File['Theme/Topic'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]
    # DF15 = XL_File['File_Name/DirName'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]
    # # DF15df = pd.DataFrame(DF15)
    # DF16 = XL_File['Date_of_modified'].loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]


    # print(DF13)

    # DF1_1 = DF11.to_string(index=False).rjust(10)
    # DF1_2 = DF12.to_string(index=False).rjust(10)
    # if DF13.empty:
    #     DF1_3 = 'No sentences have been extracted from the Account:', MSA_Name, 'and the Theme:', Theme_Name
    # else:
    #     DF1_3 = DF13.to_string(index=False).rjust(10)
    # DF1_4 = DF14.to_string(index=False).rjust(10)
    # DF1_5 = DF15.to_string(index=False).rjust(10)
    # DF1_6 = DF16.to_string(index=False).rjust(10)

    return render(request, 'CNXProtect_Utl/indexing4.html',
                  {'Account': Account, 'DF17_1': DF17_1})
                      # , 'DF1_1': DF1_1, 'DF1_2': DF1_2, 'DF1_3': DF1_3, 'DF1_4': DF1_4, 'DF1_5': DF1_5, 'DF1_6': DF1_6})


def indexing_data5(request):
    # return render(request, 'CNXProtect_Utl/indexing5.html',{'DF17_3': 700})
    if request.method == "POST":
        notesfile_1 = request.FILES['notesfile_1']
        p = notesfile_1
        # p1 = file.notesfile
        a = Master_Data(notesfile_1=notesfile_1)
        a.save()
    p3 = Master_Data.objects.values('notesfile_1')
    # print(Master_Data.objects.filter(notesfile_1__startwith = '0' ))
    p4 = pd.DataFrame(p3, columns=['notesfile_1'])
    print(p4)
    F_Len = len(p4.iloc[:]['notesfile_1'])
    print(F_Len)
    F_Count = p4.iloc[F_Len - 1]['notesfile_1']
    F_Count1 = p4.iloc[F_Len - 2]['notesfile_1']
    X_File = "C://Users/sumithra.r/Downloads/deployment/media/" + F_Count
    File_Tracker = "C://Users/sumithra.r/Downloads/deployment/media/" + F_Count1
    print("The output Master File is", X_File)
    print("The output File Tracker is", File_Tracker)
    XL_File = pd.read_excel(X_File)
    FileTracker = pd.read_excel(File_Tracker)
    # XL_File['Date_of_modified'] = pd.to_datetime(XL_File['Date_of_modified'])
    MSA_Name = request.GET.getlist('MSA_1')
    print(MSA_Name)
    Theme_Name = request.GET.getlist('Theme')

    DF17 = XL_File[XL_File['MSA'].isin(MSA_Name) & XL_File['Word'].isin(Theme_Name)]
    # DF17 = XL_File.loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]
    excel_filename_1 = "Indexing_report_" + str(datetime.datetime.today().date()) + ".xlsx"
    DF17.to_excel(BASE_DIR + '/media/' + excel_filename_1)
    DF17_1 = '/media/' + excel_filename_1

    # print(DF17_2)

    Sent = XL_File['Sentence']
    MSA = XL_File['MSA']
    Themes = XL_File['Theme/Topic']
    Keyword = XL_File['Word']
    File = XL_File['File_Name/DirName']
    print(Sent)

    Sent1 = Sent.iloc[101]
    Sent2 = Sent.iloc[102]
    Sent3 = Sent.iloc[103]
    Sent4 = Sent.iloc[104]
    Sent5 = Sent.iloc[105]
    Sent6 = Sent.iloc[106]
    Sent7 = Sent.iloc[107]
    Sent8 = Sent.iloc[108]
    Sent9 = Sent.iloc[109]
    Sent10 = Sent.iloc[110]
    Sent11 = Sent.iloc[111]
    Sent12 = Sent.iloc[112]

    MSA1 = MSA.iloc[101]
    MSA2 = MSA.iloc[102]
    MSA3 = MSA.iloc[103]
    MSA4 = MSA.iloc[104]
    MSA5 = MSA.iloc[105]
    MSA6 = MSA.iloc[106]
    MSA7 = MSA.iloc[107]
    MSA8 = MSA.iloc[108]
    MSA9 = MSA.iloc[109]
    MSA10 = MSA.iloc[110]
    MSA11 = MSA.iloc[111]
    MSA12 = MSA.iloc[112]

    Themes1 = Themes.iloc[101]
    Themes2 = Themes.iloc[102]
    Themes3 = Themes.iloc[103]
    Themes4 = Themes.iloc[104]
    Themes5 = Themes.iloc[105]
    Themes6 = Themes.iloc[106]
    Themes7 = Themes.iloc[107]
    Themes8 = Themes.iloc[108]
    Themes9 = Themes.iloc[109]
    Themes10 = Themes.iloc[110]
    Themes11 = Themes.iloc[111]
    Themes12 = Themes.iloc[112]

    File1 = File.iloc[101][78:]
    File2 = File.iloc[102][78:]
    File3 = File.iloc[103][78:]
    File4 = File.iloc[104][78:]
    File5 = File.iloc[105][78:]
    File6 = File.iloc[106][78:]
    File7 = File.iloc[107][78:]
    File8 = File.iloc[108][78:]
    File9 = File.iloc[109][78:]
    File10 = File.iloc[110][78:]
    File11 = File.iloc[111][78:]
    File12 = File.iloc[112][78:]
    File13 = File.iloc[113][78:]

    # print(pd.DataFrame(MSA_Name))
    Account = len(XL_File['MSA'].unique())
    Theme = len(XL_File['Theme/Topic'].unique())
    keywords = len(XL_File['Word'].unique())
    File_No = len(FileTracker['File_Name'].unique())
    DF17_2 = len(XL_File)
    XL = XL_File
    print(XL)
    excel_filename_3 = "Contract_Master_Data" + str(datetime.datetime.today().date()) + ".xlsx"
    XL_File.to_excel(BASE_DIR + '/media/' + excel_filename_3)
    df_filepath_Exe1 = '/media/' + excel_filename_3

# ************************************************ FROM HERE FOR FMEA *****************************************************************
    if request.method == "POST":
        notesfile_11 = request.FILES['notesfile_11']
        p = notesfile_11
        # p1 = file.notesfile
        a = Master_Data1(notesfile_11=notesfile_11)
        a.save()
    p5 = Master_Data1.objects.values('notesfile_11')
    # print(Master_Data.objects.filter(notesfile_1__startwith = '0' ))
    p6 = pd.DataFrame(p5, columns=['notesfile_11'])
    print(p6)
    F_Len = len(p6.iloc[:]['notesfile_11'])
    print(F_Len)
    F_Count_1 = p6.iloc[F_Len - 1]['notesfile_11']
    F_Count1_1 = p6.iloc[F_Len - 2]['notesfile_11']
    X_File_1 = "C://Users/sumithra.r/Downloads/deployment/media/" + F_Count_1
    File_Tracker_1 = "C://Users/sumithra.r/Downloads/deployment/media/" + F_Count1_1
    print("The output Master File is", X_File_1)
    print("The output File Tracker is", File_Tracker_1)
    XL_File_1 = pd.read_excel(X_File_1)
    FileTracker_1 = pd.read_excel(File_Tracker_1)
    # XL_File['Date_of_modified'] = pd.to_datetime(XL_File['Date_of_modified'])
    MSA_Name_1 = request.GET.getlist('Geo')
    print(MSA_Name_1)
    Theme_Name_1 = request.GET.getlist('Theme/Topic')
    print(Theme_Name_1)
    DF17_1 = XL_File_1[XL_File_1['Geo'].isin(MSA_Name_1) & XL_File_1['Key'].isin(Theme_Name_1)]
    # DF17 = XL_File.loc[(XL_File['MSA'] == MSA_Name) & (XL_File['Theme/Topic'] == Theme_Name) & (XL_File['Date_of_modified'] >= Time_in) & (XL_File['Date_of_modified'] <= Time_out)]
    excel_filename_1_1 = "Indexing_report_FMEA_" + str(datetime.datetime.today().date()) + ".xlsx"
    DF17.to_excel(BASE_DIR + '/media/' + excel_filename_1_1)
    DF17_1_1 = '/media/' + excel_filename_1
    DF17_2_1 = len(XL_File_1)
    print(DF17_2_1)

    Sent_1 = XL_File_1['Sentance']
    MSA_1 = XL_File_1['MSA']
    Themes_1 = XL_File_1['Theme/Topic']
    Keyword_1 = XL_File_1['Key']
    File_1 = XL_File_1['FileName']
    print(Sent_1)

    Account_1 = len(XL_File_1['Geo'].unique())
    Theme_1 = len(XL_File_1['Theme/Topic'].unique())
    keywords_1 = len(XL_File_1['Key'].unique())
    File_No_1 = len(FileTracker_1['FileName'].unique())
    Records_1 = len(FileTracker_1['FileName'].unique())
    DF17_2_1 = len(XL_File_1)

    Sent_1 = XL_File_1['Sentance']
    MSA_1 = XL_File_1['MSA']
    Themes_1 = XL_File_1['Theme/Topic']
    Keyword_1 = XL_File_1['Key']
    File_1 = XL_File_1['FileName']
    print(Sent)

    Sent1_1 = Sent_1.iloc[1]
    Sent2_1 = Sent_1.iloc[2]
    Sent3_1 = Sent_1.iloc[3]
    Sent4_1 = Sent_1.iloc[4]
    Sent5_1 = Sent_1.iloc[5]
    Sent6_1 = Sent_1.iloc[6]
    Sent7_1 = Sent_1.iloc[7]
    Sent8_1 = Sent_1.iloc[8]
    Sent9_1 = Sent_1.iloc[9]
    Sent10_1 = Sent_1.iloc[10]
    Sent11_1 = Sent_1.iloc[11]
    Sent12_1 = Sent_1.iloc[12]

    MSA1_1 = MSA_1.iloc[1]
    MSA2_1 = MSA_1.iloc[2]
    MSA3_1 = MSA_1.iloc[3]
    MSA4_1 = MSA_1.iloc[4]
    MSA5_1 = MSA_1.iloc[5]
    MSA6_1 = MSA_1.iloc[6]
    MSA7_1 = MSA_1.iloc[7]
    MSA8_1 = MSA_1.iloc[8]
    MSA9_1 = MSA_1.iloc[9]
    MSA10_1 = MSA_1.iloc[10]
    MSA11_1 = MSA_1.iloc[11]
    MSA12_1 = MSA_1.iloc[12]

    Themes1_1 = Themes_1.iloc[1]
    Themes2_1 = Themes_1.iloc[2]
    Themes3_1 = Themes_1.iloc[3]
    Themes4_1 = Themes_1.iloc[4]
    Themes5_1 = Themes_1.iloc[5]
    Themes6_1 = Themes_1.iloc[6]
    Themes7_1 = Themes_1.iloc[7]
    Themes8_1 = Themes_1.iloc[8]
    Themes9_1 = Themes_1.iloc[9]
    Themes10_1 = Themes_1.iloc[10]
    Themes11_1 = Themes_1.iloc[11]
    Themes12_1 = Themes_1.iloc[12]

    File1_1 = File_1.iloc[1][70:]
    File2_1 = File_1.iloc[2][70:]
    File3_1 = File_1.iloc[3][70:]
    File4_1 = File_1.iloc[4][70:]
    File5_1 = File_1.iloc[5][70:]
    File6_1 = File_1.iloc[6][70:]
    File7_1 = File_1.iloc[7][70:]
    File8_1 = File_1.iloc[8][70:]
    File9_1 = File_1.iloc[9][70:]
    File10_1 = File_1.iloc[10][70:]
    File11_1 = File_1.iloc[11][70:]
    File12_1 = File_1.iloc[12][70:]
    File13_1 = File_1.iloc[13][70:]
    XL_1 = XL_File_1
    print(XL_1)
    excel_filename_4 = "FMEA_Master_Data" + str(datetime.datetime.today().date()) + ".xlsx"
    XL_File_1.to_excel(BASE_DIR + '/media/' + excel_filename_4)
    df_filepath_Exe1_1 = '/media/' + excel_filename_4
    return render(request, 'CNXProtect_Utl/indexing5.html', {'Sent': Sent, 'MSA': MSA, 'Themes': Themes, 'Keyword': Keyword,
                                                             'File': File,'Sent_1': Sent_1, 'MSA_1': MSA_1, 'Themes_1': Themes_1, 'Keyword_1': Keyword_1,
                                                             'File_1': File_1, 'df_filepath_Exe1': df_filepath_Exe1, 'Account': Account,
                                                             'DF17_1': DF17_1, 'DF17_2': DF17_2, 'DF17_3': 700, 'Theme': Theme,
                                                             'keywords': keywords, 'File_No': File_No,
                                                             'Account_1': Account_1,
                                                             # 'DF17_1': DF17_1,  'DF17_3': 700,
                                                             'Records_1': Records_1,'Theme_1': Theme_1,
                                                             'keywords_1': keywords_1, 'File_No_1': File_No_1,'df_filepath_Exe1_1':df_filepath_Exe1_1,
                                                             'XL_1': XL_1, 'DF17_2_1': DF17_2_1,

                                                             'Sent1_1': Sent1_1, 'Sent2_1': Sent2_1, 'Sent3_1': Sent3_1, 'Sent4_1': Sent4_1, 'Sent5_1': Sent5_1,
                                                             'Sent6_1': Sent6_1, 'Sent7_1': Sent7_1, 'Sent8_1': Sent8_1, 'Sent9_1': Sent9_1, 'Sent10_1': Sent10_1,
                                                             'Sent11_1': Sent11_1, 'Sent12_1': Sent12_1,
                                                             'File1_1': File1_1, 'File2_1': File2_1,  'File3_1': File3_1,  'File4_1': File4_1,  'File5_1': File5_1,
                                                             'File6_1': File6_1,  'File7_1': File7_1,  'File8_1': File8_1,  'File9_1': File9_1,  'File10_1': File10_1,
                                                             'File11_1': File11_1,  'File12_1': File12_1,
                                                             'Themes1_1': Themes1_1,  'Themes2_1': Themes2_1,  'Themes3_1': Themes3_1,  'Themes4_1': Themes4_1, 'Themes5_1': Themes5_1,
                                                             'Themes7_1': Themes7_1,  'Themes8_1': Themes8_1, 'Themes9_1': Themes9_1,  'Themes10_1': Themes10_1,  'Themes11_1': Themes11_1,
                                                             'Themes12_1': Themes12_1,  'Themes6_1': Themes6_1,
                                                             'MSA1_1': MSA1_1, 'MSA2_1': MSA2_1, 'MSA3_1': MSA3_1,'MSA4_1': MSA4_1,'MSA5_1': MSA5_1,'MSA6_1': MSA6_1,'MSA7_1': MSA7_1,'MSA8_1': MSA8_1,
                                                             'MSA9_1': MSA9_1, 'MSA10_1': MSA10_1, 'MSA11_1': MSA11_1, 'MSA12_1': MSA12_1,

                                                             'Sent1': Sent1, 'Sent2': Sent2, 'Sent3': Sent3, 'Sent4': Sent4, 'Sent5': Sent5,
                                                             'Sent6': Sent6, 'Sent7': Sent7, 'Sent8': Sent8, 'Sent9': Sent9, 'Sent10': Sent10,
                                                             'Sent11': Sent11, 'Sent12': Sent12,
                                                             'File1': File1, 'File2': File2,  'File3': File3,  'File4': File4,  'File5': File5,
                                                             'File6': File6,  'File7': File7,  'File8': File8,  'File9': File9,  'File10': File10,
                                                             'File11': File11,  'File12': File12,
                                                             'Themes1': Themes1,  'Themes2': Themes2,  'Themes3': Themes3,  'Themes4': Themes4, 'Themes5': Themes5,
                                                             'Themes7': Themes7,  'Themes8': Themes8, 'Themes9': Themes9,  'Themes10': Themes10,  'Themes11': Themes11,
                                                             'Themes12': Themes12,  'Themes6': Themes6,
                                                             'MSA1': MSA1, 'MSA2': MSA2, 'MSA3': MSA3,'MSA4': MSA4,'MSA5': MSA5,'MSA6': MSA6,'MSA7': MSA7,'MSA8': MSA8,
                                                             'MSA9': MSA9, 'MSA10': MSA10, 'MSA11': MSA11, 'MSA12': MSA12})
                      # , 'DF1_1': DF1_1, 'DF1_2': DF1_2, 'DF1_3': DF1_3, 'DF1_4': DF1_4, 'DF1_5': DF1_5, 'DF1_6': DF1_6})


def score_card(request):
    from tabula import read_pdf
    from tabulate import tabulate
    import pandas as pd
    import numpy as np
    # reads table from pdf file
    df = read_pdf("C://Users/RNALAB/OneDrive - Concentrix Corporation/Score_Card/Scorecard.pdf",
                  pages="all")  # address of pdf file
    for sht in range(0, len(df)):
        if (len(df[sht]) > 10):
            file = df[sht]
            # if (df1.columns.size>10):
            print(file.loc[0].values[1])
            if file.loc[0].values[1] == 'nan':
                Column_list = file.loc[0, :]
                Column_header = Column_list.values

            elif file.loc[1].values[1] != 'nan':
                Column_list = file.loc[1, :]
                Column_header = Column_list.values

            if (df[sht].columns.size) > 10:
                str_sz = [10, 14]
                for str_sz_1 in range(0, len(str_sz)):
                    val = str_sz[str_sz_1]
                    # if len(Column_header[val]) > 20:

                    df2 = file.iloc[:, val].str.split(expand=True)
                    file = pd.concat([file, df2], axis=1)

                    df1 = file.reset_index(drop=True)
                    df2 = file.reset_index(drop=True)
                    # both 2 ways to get the same result
                    # df4 = file.join(df2)

            file.to_excel(
                'C:\\Users\RNALAB\OneDrive - Concentrix Corporation\Score_Card\Score_card_merged1_6.58_Table_' + str(
                    sht + 1) + '.xlsx', sheet_name='Sheet' + str(sht + 1))
            # df3.save()
            # , index = False)
            # writer = pd.ExcelWriter('C:\\Users\sumithra.r\OneDrive - Concentrix Corporation\Score_Card\Score_card_merged1_6.58.xlsx.xlsx', engine='xlsxwriter')
        else:
            file.to_excel(
                'C:\\Users\RNALAB\OneDrive - Concentrix Corporation\Score_Card\Score_card_merged1_6.58_Table_' + str(
                    sht + 1) + '.xlsx', sheet_name='Sheet' + str(sht + 1))
    excel_filename_score = "Score_card" + str(datetime.datetime.today().date()) + ".xlsx"
    file.to_excel(BASE_DIR + '/media/' + excel_filename_score)
    score = '/media/' + excel_filename_score

    return render(request, 'CNXProtect_Utl/score_card.html', {'score': score})


def toast(request):
    messages.success(request, "Will check with the existing keywords")
# def insert_Master_keyword(request):
#     p3 = Master_Data.objects.values('notesfile_1')
#     # print(Master_Data.objects.filter(notesfile_1__startwith = '0' ))
#     p4 = pd.DataFrame(p3, columns=['notesfile_1'])
#     print(p4)
#     F_Len = len(p4.iloc[:]['notesfile_1'])
#     print(F_Len)
#     F_Count = p4.iloc[F_Len-1]['notesfile_1']
#     X_File = "C://Users/RNALAB/Concentrix Corporation/CNX_Protect/media/" + F_Count
#     print(X_File)
#     XL_File = pd.read_excel(X_File)
#     # print("Column headings:", XL_File.columns)
#     # print(XL_File.columns)
#     # df = pd.DataFrame(XL_File)
#     # U_account = (df['MSA'].unique())
#     Account = XL_File['MSA'].unique()
#     return render(request, 'CNXProtect_Utl/Contract_Library.html', {'Account': Account})


def indexing(request):
    p3 = Master_Data.objects.values('notesfile_1')
    # print(Master_Data.objects.filter(notesfile_1__startwith = '0' ))
    p4 = pd.DataFrame(p3, columns=['notesfile_1'])
    print(p4)
    F_Len = len(p4.iloc[:]['notesfile_1'])
    print(F_Len)
    F_Count = p4.iloc[F_Len - 1]['notesfile_1']
    X_File = "C://Users/sumithra.r/Downloads/deployment/media/" + F_Count
    print(X_File)
    XL_File = pd.read_excel(X_File)
    # print("Column headings:", XL_File.columns)
    # print(XL_File.columns)
    # df = pd.DataFrame(XL_File)
    # U_account = (df['MSA'].unique())
    # Account = XL_File['MSA'].unique()

    DF11 = XL_File['FileName'].loc[XL_File['MSA'] == 'RSL']
    DF1_2 = XL_File['Key'].loc[XL_File['MSA'] == 'RSL']
    DF1_3 = XL_File['Sentance'].loc[XL_File['MSA'] == 'RSL']
    DF1_4 = XL_File['MSA'].loc[XL_File['MSA'] == 'RSL']

    # print(Slct_data)
    # Slct_data = XL_File.loc[XL_File['MSA'] == Account['Americanexpress']]

    # print(Account)
    # for acc in range(len(Account)):
    #     Slct_data = XL_File.loc[XL_File['MSA'] == Account[acc]]
    #     print(Slct_data)
    return render(request, 'CNXProtect_Utl/Contract_Library.html',
                  {'DF1_1': DF11, 'DF1_2': DF1_2, 'DF1_3': DF1_3, 'DF1_4': DF1_4})


def countries_view(request):
    if request.method == 'POST':
        form = CountryForm(request.POST)
        if form.is_valid():
            countries = form.cleaned_data.get('countries')
    else:
        form = CountryForm
    return render(request, "CNXProtect_Utl/indexing1.html", {'form': form},
                  context_instance == RequestContext(request))

    # if df.columns == 'MSA':
    #     print(XL_File.MSA.unique())

    # KW1 = pd.DataFrame(KW2, columns=['keywords'])
    # print(p)
    # p = File.objects.values('notesfile')
    # print(p)
    # print("C:\Users\RNALAB\Concentrix Corporation\CNX_Protect\media\media" + p)
    # messages.success(request, 'Master File submitted successfully!!')
    # return render(request, 'CNXProtect_Utl/keyword_list.html', p)
    #     return redirect('Contract_Library')
    # else:
    #     messages.error(request, 'Master Files not submitted successfully!!')
    #     return redirect('Contract_Library')
    # File_excel = open('C:\\Users\RNALAB\\Concentrix Corporation\\Utility\\master_data_exp.xlsx')
    #
    # , {'df_filepath_out1_1': df_filepath_out1_1})
# def cl_files(request, context=None):
#     global Statement1
#     import json
#     # global df, PP6, q, Statement, Frq, No_of_word, No_of_Line
#
#     def split_into_sentences(text):
#         text = " " + text + "  "
#         R = dec.findall(text)
#         S = [j[0] + '<decimal>' + j[1] for j in [dig.findall(i) for i in R]]
#         for i in range(len(R)):
#             text = text.replace(R[i], S[i])
#         text = text.replace("\n", " ")
#         text = text.replace("Œ", "<stop>")
#         text = re.sub(prefixes, "\\1<prd>", text)
#         text = re.sub(websites, "<prd>\\1", text)
#         if "Ph.D" in text: text = text.replace("Ph.D.", "Ph<prd>D<prd>")
#         text = re.sub("\s" + alphabets + "[.] ", " \\1<prd> ", text)
#         text = re.sub(acronyms + " " + starters, "\\1<stop> \\2", text)
#         text = re.sub(alphabets + "[.]" + alphabets + "[.]" + alphabets + "[.]", "\\1<prd>\\2<prd>\\3<prd>", text)
#         text = re.sub(alphabets + "[.]" + alphabets + "[.]", "\\1<prd>\\2<prd>", text)
#         text = re.sub(" " + suffixes + "[.] " + starters, " \\1<stop> \\2", text)
#         text = re.sub(" " + suffixes + "[.]", " \\1<prd>", text)
#         text = re.sub(" " + alphabets + "[.]", " \\1<prd>", text)
#         if "”" in text: text = text.replace(".”", "”.")
#         if "\"" in text: text = text.replace(".\"", "\".")
#         if "!" in text: text = text.replace("!\"", "\"!")
#         if "?" in text: text = text.replace("?\"", "\"?")
#         text = text.replace(".", ".<stop>")
#         text = text.replace("?", "?<stop>")
#         text = text.replace("!", "!<stop>")
#         text = text.replace("<prd>", ".")
#         sentences = text.split("<stop>")
#         # sentences = sentences[:-1]
#         sentences = [s.strip() for s in sentences]
#         return sentences
#
#     def readtxt(filename):
#         doc = docx.Document(filename)
#         fullText = []
#         for para in doc.paragraphs:
#             fullText.append(para.text)
#         return '\n'.join(fullText)
#
#     # _______________________________________________
#     #     WordSearch / PATTERN SEARCH FROM SENTENCES
#     def search(word, sentences):
#         return [i for i in sentences if re.search(r'\b%s\b' % word, i)]
#
#     # _______________________________________________
#     #
#     # This function is taking Key and dataframe and returning all the rows and columns with matched Keywords
#     def xsearch(regex: str, df, case=False):
#         """Search all the text columns of `df`, return rows with any matches."""
#         textlikes = df.select_dtypes(include=[object, "string"])
#         return df[
#             textlikes.apply(
#                 lambda column: column.str.contains(regex, regex=True, case=case, na=False)
#             ).any(axis=1)]
#
#     # _______________________________________________
#     # This function is searching the Key inside the filtered list
#     def indexHaveSubstring(lst, substring):
#         ls = list()
#         for i in range(len(lst)):
#             subs = re.compile(Ke).search(str(lst[i]))
#             if subs != None:
#                 ls.append(i)
#         return ls
#
#     # _______________________________________________
#     #####     DATE SEARCH PATTERNS
#     import calendar
#     dat = re.compile(
#         r'(?:\d{1,2}(?:(?:-|/)|(?:th|st|nd|rd)?\s))?(?:(?:(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)(?:(?:-|/)|(?:,|\.)?\s)?)?(?:\d{1,2}(?:(?:-|/)|(?:th|st|nd|rd)?\s))?)(?:\d{2,4})')
#     dat = r'(?:\d{1,2}[-/th|st|nd|rd\s]*)?(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)|jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)[a-z\s,.]*(?:\d{1,2}[-/th|st|nd|rd)\s,]*)+(?:\d{3})+'
#     dat = r'(?:\d{1,2}[-/th|st|nd|rd\s]*|[-/-])?(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)|jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)|\d{1,2}?)[a-z\s,.]*(?:\d{1,2}[-/th|st|nd|rd)\s,]*)+(?:\d{3})+'
#     full_months = [month for month in calendar.month_name if month]
#     short_months = [d[:3] for d in full_months]
#     months = '|'.join(short_months + full_months)
#     sep = r'[.,]?\s+'  # seperators
#     day = r'\d+'
#     year = r'\d+'
#     day_or_year = r'\d+(?:\w+)?'
#
#     KW2 = Dictionary.objects.values('keywords')
#     KW1 = pd.DataFrame(KW2, columns=['keywords'])
#     # KW = KW1.at[:,'keywords']
#     Dictionary_Word = KW1
#     print(Dictionary_Word)
#     # print(KW1)
#
#     flg = "lower"
#     DF = pd.DataFrame(
#         columns=["FileName/DirName", "Word", "Theme/Topic", "Sentence", "SentenceNo",
#                  "PageNo", "Text Summerization", "Frequency of Keyword",
#                  "No_of_word in the paragraph", "No_of_Line in the parapraph"])
#     # print(DF)
#     df1 = pd.DataFrame()
#     FileTrack = pd.DataFrame(columns=["FileName", "Page", "Status", "CharCount"])
#     # excel_filename2 = "Contract_File_output" + str(datetime.datetime.today().date()) + ".xlsx"
#     # p = "C:\\Users\RNALAB\Concentrix Corporation\Contract Lists - Contract Library\Broadcom (CA Inc.)\CA Inc-CNX US(MSA)07-24-20exe.pdf"
#     p1 = File.objects.values('notesfile')
#     p2 = pd.DataFrame(p1, columns=['notesfile'])
#     # KW = KW1.at[:,'keywords']
#     p3 = p2.iloc[-1]['notesfile']
#     # print(p3)
#     p = "C://Users/RNALAB/Concentrix Corporation/CNX_Protect/media/" + p3
#     # BASE_DIR1 = os.path.join(BASE_DIR, 'media')
#     # p = BASE_DIR1+p3
#     print(p)
#     # print(p)
#     pdfFileObj = open(p, 'rb')
#     plen = 0
#     p1 = p
#     pdfReader = PyPDF4.PdfFileReader(pdfFileObj,
#                                      strict=False)  ###   STRICT + FALSE ADDED TO REMOVE THE ERROR CHECK FOR ACCURACY
#     if not pdfReader.isEncrypted:
#         for k in range(0, Dictionary_Word.shape[1]):
#             for j in range(0, Dictionary_Word.shape[0]):
#                 Keyword_present = 1;
#                 if str(Dictionary_Word.iloc[j, k]) != 'nan':
#                     # for(m in splitted){
#                     # print(names(Dictionary[k]))
#                     #
#                     q = Dictionary_Word.iloc[j, k]
#                     print(q)
#                     tXt = "";
#                     PGtXt = ""
#                     for PG in range(pdfReader.numPages):
#                         PGNo = PG
#                         print(PGNo)
#                         pageObj = pdfReader.getPage(PG)
#                         PGtXt = pageObj.extractText()
#                         PGtXt = PGtXt.replace(
#                             ' !"#$%&\'()*+,-./012-345627\n3,&"8%))9"#:;3<%&$9=%)35\'&%%>%#?\n3,:\'%\n3@3',
#                             '').replace('!%3', '').replace('!&3', '').replace('!$3', '').replace('!#3', '').replace(
#                             '!"3', '').replace('!!3', '').replace('!*3', '').replace('!$3',
#                                                                                      '').replace('!*3', '').replace(
#                             '!3', '').replace('#3', '').replace('$3', '').replace('%3', '').replace("'3", '').replace(
#                             '(3', '').replace(')3', '').replace('&3', '')
#                         tXt = tXt + '\n-\r-\n' + PGtXt
#                         tXt = tXt.lower()
#                         PGtXt = PGtXt.strip('\n');
#                         PGtXt = PGtXt.replace('\n', '');
#                         PGtXt = PGtXt.strip('\t');
#                         PGtXt = PGtXt.replace('\t', ' ');
#                         PGtXt = re.sub("\s+", ' ', PGtXt)
#                         plen = plen + len(PGtXt)
#                         PGt = split_into_sentences(PGtXt)
#                         if (flg == "lower"):
#                             patt = q.lower()
#                             PGt = [x.lower() for x in PGt]
#                         else:
#                             patt = q;
#                         ### tolower
#                         word = patt
#                         # print(q) #                    Test
#                         if '+DATE' in q:
#                             word = q.replace('+DATE', '')
#                             word = word.lower()
#                             PG1 = [[(word in x), PGt.index(x)] for x in PGt]
#                             PG2 = list(compress(PGt, [item[0] for item in PG1]))
#                             PG3 = [item[1] for item in PG1 if item[0] == True]
#                             PP1 = [[bool(re.search(dat, x)), PG2.index(x)] for x in PG2]
#                             PP2 = list(compress(PP1, [item[0] for item in PP1]))
#                             PP3 = [item[1] for item in PP1 if item[0] == True]
#                             PP4 = [PG2[item] for item in PP3]
#                         elif '+' in patt:
#                             word = word.split(sep='+')
#                             dat = word[1]
#                             PG1 = [[(word[0] in x), PGt.index(x)] for x in PGt]
#                             PG2 = list(compress(PGt, [item[0] for item in PG1]))
#                             PG3 = [item[1] for item in PG1 if item[0] == True]
#                             PP1 = [[bool(dat in x), PG2.index(x)] for x in PG2]
#                             PP2 = list(compress(PP1, [item[0] for item in PP1]))
#                             PP3 = [item[1] for item in PP1 if item[0] == True]
#                             PP4 = [PG2[item] for item in PP3]
#                         else:
#                             PG1 = [[((' ' + word + ' ' in x) or (word + ' ' in x) or (word + '.' in x)), PGt.index(x)]
#                                    for x in PGt]
#                             # PG2 = list(compress(PGt,[item[0] for item in PG1]))
#                             PG2 = (compress(PGt, [item[0] for item in PG1]))
#                             PG3 = [item[1] for item in PG1 if item[0] == True]
#                             PP3 = PG3
#                             import numpy as np
#                             if PP3 != []:
#                                 Pp3 = (np.array(PP3));
#                                 for lst in range(len(Pp3)):
#                                     PP5 = Pp3[lst];
#                                     if (len(PGt[PP5]) < 150) and (PP5 < len(PGt) - 1):
#                                         # Sentenses are:
#                                         PP6 = PGt[PP5 - 1] + PGt[PP5] + PGt[PP5 + 1]
#                                     else:
#                                         PP6 = PGt[PP5]
#                                     Keyword_present += Keyword_present;
#                                     # print('keyword:', q,'\n')
#
#                                     # print("Sentence:", PP6,'\n')
#                                     # text analysis starts here:PP
#                                     Main_sentance = []
#                                     sents = nltk.sent_tokenize(PP6.lower())
#                                     # emails = re.findall(r"[a-z0-9\.\-+_]+@[a-z0-9\.\-+_]+\.[a-z]+", text)
#                                     # print emails
#                                     No_of_word = len(sents[0])
#                                     No_of_Line = int(No_of_word / 120)
#                                     lemmatizer = WordNetLemmatizer()
#                                     for i in range(len(sents)):
#                                         emails = re.findall(r"[a-z0-9\.\-+_]+@[a-z0-9\.\-+_]+\.[a-z]+", sents[i])
#                                         print(emails)
#                                         if len(emails) != 0:
#                                             Ph_number = re.findall(
#                                                 r"((?:\+\d{2}[-\.\s]??|\d{4}[-\.\s]??)?(?:\d{3}[-\.\s]??\d{3}[-\.\s]??\d{4}|\(\d{3}\)\s*\d{3}[-\.\s]??\d{4}|\d{3}[-\.\s]??\d{4}))",
#                                                 sents[i])
#                                             words = nltk.word_tokenize(sents[i])
#                                             POS = nltk.pos_tag(words)
#                                             print('POS tagging:', POS, '\n')
#                                             Tagged_words = []
#                                             for words, tag in POS:
#                                                 if (tag == 'CD' or tag == 'NNS'):
#                                                     Tagged_words.append(words)
#                                             print('Tagged_words:', Tagged_words, '\n')
#
#                                             if (".pdf" in p and len(tXt) < 150):
#                                                 ImagePDF = "Image PDF"
#                                             else:
#                                                 ImagePDF = 'Proper PDF File'
#                                             df = pd.Series(
#                                                 [p1, ImagePDF, q, Dictionary.columns[k], PP6, PP5, 0, PG + 1, emails,
#                                                  Ph_number, Tagged_words],
#                                                 index=["FileName/DirName", "ImagePDF", "Word", "Theme/Topic",
#                                                        "Sentence", "SentenceNo", "TotalCount", "PageNo", "email ID",
#                                                        "Phone Number", "Timeline"])
#                                             df1 = df1.append(df, ignore_index=True)
#                                             # print("Keyword:", q)
#                                             # print("Sentence:", PP6)
#                                 # DF = DF.append(df1)
#                                 # df1 = pd.DataFrame()
#                             df2 = pd.Series([p1, "File", "Read", plen],
#                                             index=["FileName", "Page", "Status", "CharCount"])
#                             FileTrack = FileTrack.append(df2, ignore_index=True)
#         if Keyword_present <= 1:
#             PP6 = "Keyword not present"
#             PP5 = "Nil"
#             PG = "Nil"
#             Txt_Anlys = "Nil"
#             Statement = "Keyword not present"
#             Frq = "Nil"
#             No_of_word = "Nil"
#             No_of_Line = "Nil"
#             emails = "Nil"
#             Ph_number = "Nil"
#             Tagged_words = "Nil"
#             df = pd.Series(
#                 [p1, q, Dictionary_Word.columns[k], PP6, PP5, PG, Txt_Anlys, Frq, No_of_word, No_of_Line],
#                 index=["FileName/DirName", "ImagePDF", "Word", "Theme/Topic",
#                        "Sentence", "SentenceNo", "TotalCount", "PageNo", "email ID",
#                        "Phone Number", "Timeline"])
#             df1 = df1.append(df, ignore_index=True)
#
#     DF = DF.append(df1)
#     df2 = pd.Series([p3, "File", "Read", plen], index=["FileName", "Page", "Status", "CharCount"])
#     FileTrack = FileTrack.append(df2, ignore_index=True)
#     # print(FileTrack)
#     # excel_filename = "Contract_File_Search" + str(p1) + "_" + str(datetime.datetime.today().date()) + ".xlsx"
#     excel_filename1 = "Email and Phone_Tracker" + str(datetime.datetime.today().date()) + ".xlsx"
#     FileTrack.to_excel(BASE_DIR + '/media/' + excel_filename1)
#     df_filepath_trk = '/media/' + excel_filename1
#
#     excel_filename2 = "Email and Phone_output" + str(datetime.datetime.today().date()) + ".xlsx"
#     DF.to_excel(BASE_DIR + '/media/' + excel_filename2)
#     df_filepath_out = '/media/' + excel_filename2
#
#     excel_filename3 = "Contract_File_Tracker" + str(datetime.datetime.today().date()) + ".xlsx"
#     FileTrack.to_excel(BASE_DIR + '/media/' + excel_filename3)
#     df_filepath = '/media/' + excel_filename3
#     # DF_File =DF.to_excel( "C:\\Users\RNALAB\Concentrix Corporation\CNX_Protect\media\result.xlsx")
#
#     # return render(request, 'CNXProtect_Utl/cl_files.html', print({'result': DF}))
#     # return render(request, 'CNXProtect_Utl/cl_files.html', {'DF': DF})
#     # return render(request, 'CNXProtect_Utl/login_page.html', {'DF': DF})
#     return render(request, 'CNXProtect_Utl/Contract_Library.html',
#                   {'DF': DF, 'q': q, 'p3': p3, 'PP6': PP6, 'Statement': Statement, 'emails': emails,
#                    'Ph_number': Ph_number, 'Tagged_words': Tagged_words
#                    'excel_filename1': excel_filename1, 'Statement1': Statement1,
#                    'Keyword': Keyword, 'Txt_Anlys1': Txt_Anlys1,
#                    'df_filepath_trk': df_filepath_trk, 'df_filepath_out': df_filepath_out})

# def list(request):
#     # Handle file upload
#     if request.method == 'POST':
#         form = DocumentForm(request.POST, request.FILES)
#         if form.is_valid():
#             newdoc = Document(docfile = request.FILES['docfile'])
#             newdoc.save()
#
#             # Redirect to the document list after POST
#             return HttpResponseRedirect(reverse('myapp.views.list'))
#     else:
#         form = DocumentForm() # A empty, unbound form
#
#     # Load documents for the list page
#     documents = Document.objects.all()
#
#     # Render list page with the documents and the form
#     return render('keyword_list.html',{'documents': documents, 'form': form}, context_instance=RequestContext(request) )

# def customer(request, pk_test):
#     customer = Customer.objects.get(id=pk_test)
# orders = customer.order_set.all()
# order_count = orders.count()
#
# myFilter = OrderFilter(request.GET, queryset=orders)
# orders = myFilter.qs
#
# context = {'customer': customer, 'orders': orders, 'order_count': order_count,
#            'myFilter': myFilter}
# # Create your views here.
# def register(request):
#     if request.method == 'POST':
#         first_name = request.POST['first_name']
#         last_name = request.POST['last_name']
#         username = request.POST['username']
#         password1 = request.POST['password1']
#         password2 = request.POST['password2']
#         email = request.POST['email']
#
#         if password1 == password2:
#             if User.objects.filter(username=username).exists():
#                 messages.info(request, 'User Taken')
#                 return redirect('register')
#             elif User.objects.filter(email=email).exists():
#                 messages.info(request, 'Email Taken')
#                 return redirect('register')
#             else:
#                 user = User.objects.create_user(username=username, password=password1, email=email,
#                                                 first_name=first_name, last_name=last_name)
#                 user.save()
#                 messages.info(request, 'User Taken')
#                 return redirect('login')
#         else:
#             messages.info(request, 'Password not Matching')
#             return redirect('register')
#         return render(request, 'CNXProtect_Utl/login_page.html')
#     else:
#         return render(request, 'CNXProtect_Utl/register.html')
# def back_to_page(request):
#     return redirect('/login_page#tab-content-4')
# def userPage(request):
#     context = {}
#     return render(request, 'CNXProtect_Utl/index.html', context)
# Create your views here.
# def register(request):
#     if request.method == 'POST':
#         first_name = request.POST['first_name']
#         last_name = request.POST['last_name']
#         username = request.POST['username']
#         password1 = request.POST['password1']
#         password2 = request.POST['password2']
#         email = request.POST['email']
#
#         if password1 == password2:
#             if User.objects.filter(username=username).exists():
#                 messages.info(request, 'Uses Name already created!')
#                 return redirect('register')
#             elif User.objects.filter(email=email).exists():
#                 messages.info(request, 'Email already taken!')
#                 return redirect('register')
#             else:
#                 user = User.objects.create_user(username=username, password=password1, email=email,
#                                                 first_name=first_name, last_name=last_name)
#                 user.save()
#                 messages.info(request, 'User account created!')
#                 return redirect('login')
#         else:
#             messages.info(request, 'Password not Matching')
#             return redirect('register')
#         return redirect('/')
#     else:
#         return render(request, 'CNXProtect_Utl/register.html')
#
# @unauthenticated_user
# def register(request):
#     form = CreateUserForm()
#     if request.method == 'POST':
#         form = CreateUserForm(request.POST)
#         if form.is_valid():
#             user = form.save()
#             username = form.cleaned_data.get('username')
#
#             group = Group.objects.get(name='customer')
#             user.groups.add(group)
#
#             messages.success(request, 'Account was created for ' + username)
#
#             return redirect('login')
#
#     context = {'form': form}
#     return render(request, 'accounts/register.html', context)
# @login_required(login_url='login')
# @allowed_users(allowed_roles=['admin'])
# @unauthenticated_user
# def login(request):
#     if request.method == 'POST':
#         username = request.POST['username']
#         password = request.POST['password']
#
#         user = auth.authenticate(username=username, password=password)
#         if user is not None:
#             auth.login(request, user)
#             # return redirect('/')
#             return render(request, 'CNXProtect_Utl/login_page.html')
#         else:
#             messages.info(request, 'User Name or Password not matched!!!')
#             return redirect('login')
#     else:
#         messages.info(request, 'User Name or Password not matched!!!')
#         return render(request, 'CNXProtect_Utl/login.html')
# def read_file(request):
#     if request.method == "POST":
#         form = UploadFileForm(request.POST, request.FILES)
#         if form.is_valid():
#             handle_uploaded_file(request.FILES['file'])
#             context = {'msg': '<span style = "color: green;"> File Successfully uploaded </span>'}
#             return render(request, "login_page.html", context)
#     else:
#         form = UploadFileForm()
#         print(form)
#     # context = {'msg': '<span style = "color: green;"> File Successfully uploaded </span>'}
#     return render(request, 'CNXProtect_Utl/login_page.html', {'form': form})
#
#
# def handle_uploaded_file(f):
#     with open(f.name, 'wb+') as destination:
#         for chunk in f.chunks():
#             destination.write(chunk)
# extension = os.path.splitext(upload_file.name)[1]
# rename = datetime.datetime.now().strftime("%Y_%m_%d %H_%M_%S") + extension
# fss = FileSystemStorage()
# filename = fss.save(rename, upload_file)
# file1 = File(file=rename)
# file1.save()
# upload_file_path = fss.path(filename)
#         render(p, 'read_file.html', {
# #         'upload_file_path': upload_file_path
# #     })
# # else:
# #      return render(p, 'read_file.html')
# dat = re.compile(rf'(?:{day}{sep})?(?:{months}){sep}{day_or_year}(?:{sep}{year})?')
# def read_file(File1):
#     # p = "C:\\Users\RNALAB\Concentrix Corporation\Contract Lists - Contract Library\ARS-CNX US(CCA-SOW 1-CO 1)03-24-2021exe.pdf"
#     p = File1
#     #pdfFileObj = open(p, 'rb')
#     # csv_data = pd.read_csv(csv_filename)
#     return p
# def file_read(request):
#
#     return p
# def read_file(file_path):
#     p = file_path
#     return p
# def cl_files(Dictionary, keyword_id):
# def cl_files(Dictionary: {objects}, keyword_id):
# def cl_files(request, Keys):
# # KW = str(Keys)
# KW1 = Dictionary.objects.all()
# KW = str(KW1)
# KW = str(Dictionary.objects.all())

# return redirect('CNXProtect_Utl/login_page', {'active_tab': 'tab-content-4'})
# return redirect('keyword_list/')
# return render(request, 'CNXProtect_Utl/keyword_list.html')

# def simple_upload(request):
#     if request.method == 'POST':
#         person_resource = DictionaryResource()
#         dataset = Dataset()
#         new_person = request.FILES['myfile']
#         if not new_person.name.endwith('xlsx'):
#             messages.info(request,'wrong format')
#             return render(request,'upload.html')
# # from .models import CNXProtect_Utl
# def cl_files(request):
#     return render(request, 'CNXProtect_Utl/cl_files.html')

# return render(request, 'CNXProtect_Utl/cl_files.html')
# def keyword():
#     context = {'keyword_list': Dictionary.objects.all()}
#     return (context)
# def keywords_item(Keyword):
#     for Keyword in keyword_list:
#         Keys = Keyword.keywords
#     print(Keys)
#     # return render(request, 'keyword_list.html', {'Keys': Keys})
#     return Keys

# def text_data_read(textfile):
#
#     file = open(textfile, 'r')
#     text = file.read()
#     return text
